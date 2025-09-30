"""
默安SAST报告解析器实现
基于现有server.py中的解析逻辑
"""
import re
from typing import Dict, Any
from docx import Document
from .base import ISASTParser


class MoAnParser(ISASTParser):
    """默安SAST报告解析器"""

    def can_parse(self, docx_path: str) -> bool:
        """检测是否为默安报告格式"""
        try:
            doc = self._get_document(docx_path)

            # 检查默安特有的章节结构
            has_vulnerability_section = False
            has_code_standard_section = False

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if "四、漏洞详情" in text:
                    has_vulnerability_section = True
                elif "六、代码规范风险详情" in text:
                    has_code_standard_section = True

            # 默安格式：必须有漏洞详情或代码规范风险详情章节
            return has_vulnerability_section or has_code_standard_section

        except Exception:
            return False

    def get_vendor_name(self) -> str:
        """获取厂商名称"""
        return "moan"

    def get_priority(self) -> int:
        """获取解析器优先级（默安优先级较高）"""
        return 10

    def parse_docx_to_json(self, docx_path: str) -> Dict[str, Any]:
        """解析docx为标准JSON格式"""
        # 使用现有的解析逻辑
        all_docx_json = self._parse_docx_structure(docx_path)
        vulnerabilities = []

        for item in all_docx_json:
            if item["title"] in ["四、漏洞详情", "六、代码规范风险详情"]:
                vulnerabilities.extend(self._transform_json([item])[item["title"]])

        return {
            "vendor": self.get_vendor_name(),
            "report_type": "SAST",
            "vulnerabilities": vulnerabilities
        }

    def _get_paragraph_heading_level(self, paragraph):
        """获取段落标题级别"""
        if paragraph.style.name.startswith('Heading'):
            try:
                return int(paragraph.style.name.replace('Heading', ''))
            except ValueError:
                return 0
        return 0

    def _process_table(self, table):
        """处理表格数据"""
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = ""
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        cell_text += paragraph.text + " "
                row_data.append(cell_text.strip())
            table_data.append(row_data)
        return table_data

    def _get_docx_elements_in_order(self, document):
        """按顺序获取文档元素"""
        body = document._element.body
        elements = []
        for child in body.iterchildren():
            if child.tag.endswith('p'):
                for paragraph in document.paragraphs:
                    if paragraph._element is child:
                        elements.append(('paragraph', paragraph))
                        break
            elif child.tag.endswith('tbl'):
                for table in document.tables:
                    if table._element is child:
                        elements.append(('table', table))
                        break
        return elements

    def _parse_docx_structure(self, docx_path):
        """解析docx文档为JSON结构"""
        doc = self._get_document(docx_path)
        elements = self._get_docx_elements_in_order(doc)
        result = []
        heading_stack = [(-1, None, result)]
        current_content = []

        for elem_type, elem in elements:
            if elem_type == 'paragraph':
                level = self._get_paragraph_heading_level(elem)

                if level > 0:
                    if heading_stack[-1][1] is not None:
                        heading_stack[-1][1]["content"] = current_content

                    heading_node = {
                        "title": elem.text,
                        "content": []
                    }
                    current_content = []

                    while heading_stack[-1][0] >= level:
                        heading_stack.pop()

                    parent_level, parent_node, parent_container = heading_stack[-1]

                    if isinstance(parent_container, list):
                        parent_container.append(heading_node)
                    else:
                        if "children" not in parent_container:
                            parent_container["children"] = []
                        parent_container["children"].append(heading_node)

                    heading_stack.append((level, heading_node, heading_node))
                elif elem.text.strip():
                    current_content.append({"type": "text", "value": elem.text})

            elif elem_type == 'table':
                table_data = self._process_table(elem)
                current_content.append({"type": "table", "value": table_data})

        if heading_stack[-1][1] is not None:
            heading_stack[-1][1]["content"] = current_content

        return result

    def _transform_json(self, json_data):
        """转换JSON结构为漏洞报告格式"""
        result = {}
        issue_level_pattern = re.compile(r"【(低危|中危|高危|提示)】.*?漏洞数：(\d+)")

        for item in json_data:
            main_title = item["title"]
            result[main_title] = []
            if "children" not in item:
                continue

            for child in item["children"]:
                match = issue_level_pattern.search(child["title"])
                if match:
                    issue_level = match.group(1)
                    issue_count = match.group(2)

                    if issue_level == "提示":
                        issue_level = "Notice"
                    elif issue_level == "低危":
                        issue_level = "Low"
                    elif issue_level == "中危":
                        issue_level = "Medium"
                    elif issue_level == "高危":
                        issue_level = "High"

                    if issue_level not in ["Medium", "High"]:
                        continue

                    issue = {
                        "issue_title": child["title"],
                        "issue_level": issue_level,
                        "issue_count": issue_count,
                        "issue_desc": "",
                        "fix_advice": "",
                        "code_sample": "",
                        "code_list": []
                    }

                    for section in child["children"]:
                        section_title = section["title"]
                        section_content = section["content"]

                        if section_title == "漏洞描述":
                            for content_item in section_content:
                                if content_item["type"] == "text":
                                    issue["issue_desc"] = content_item["value"]

                        elif section_title == "修复建议":
                            for content_item in section_content:
                                if content_item["type"] == "text":
                                    issue["fix_advice"] = content_item["value"]

                        elif section_title == "代码示例":
                            for content_item in section_content:
                                if content_item["type"] == "text":
                                    issue["code_sample"] = content_item["value"]

                        elif re.match(r'^NO\.\d+\.\s代码位置$', section_title):
                            code_item = {}

                            if section_content and section_content[0]["type"] == "text":
                                code_location_num = section_content[0]["value"]
                                splitor = code_location_num.split(":")
                                code_item["code_location"] = splitor[0]
                                code_item["code_line_num"] = splitor[1]

                            if "children" in section:
                                for child_section in section["children"]:
                                    for content_item in child_section["content"]:
                                        if content_item["type"] == "table":
                                            code_item["code_details"] = content_item["value"][-1][1]

                            issue["code_list"].append(code_item)

                    result[main_title].append(issue)

        return result