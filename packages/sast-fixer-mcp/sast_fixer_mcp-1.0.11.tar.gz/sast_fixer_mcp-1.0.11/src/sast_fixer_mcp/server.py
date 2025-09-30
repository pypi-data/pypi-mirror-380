import logging
from pathlib import Path
from typing import Optional, List
from mcp.server import Server
from mcp.server.stdio import stdio_server
from mcp.types import (
    TextContent,
    Tool,
)
from enum import Enum
import os
import csv
import json
import re
from pydantic import BaseModel, Field
from docx import Document

# Import the new parser architecture
from .parsers import sast_parser_factory

class ConvertSastDocxToJson(BaseModel):
    file_path: str = Field(description="SAST报告docx文件路径")
    working_directory: Optional[str] = Field(default=None, description="工作目录路径，默认为当前目录")

class GetPendingVulnerabilityJsonFiles(BaseModel):
    working_directory: Optional[str] = Field(default=None, description="工作目录路径，默认为当前目录")

class GenerateCsvReport(BaseModel):
    working_directory: Optional[str] = Field(default=None, description="工作目录路径，默认为当前目录")

class SastTools(str, Enum):
    CONVERT_DOCX_TO_JSON = "convert_sast_docx_to_json"
    GET_PENDING_FILES = "get_pending_vulnerability_json_files" 
    GENERATE_CSV_REPORT = "generate_csv_report"

def get_paragraph_heading_level(paragraph):
    """获取段落标题级别"""
    if paragraph.style.name.startswith('Heading'):
        try:
            return int(paragraph.style.name.replace('Heading', ''))
        except ValueError:
            return 0
    return 0

def process_table(table):
    """处理表格数据"""
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            cell_text = ""
            for paragraph in cell.paragraphs:
                if paragraph.text.strip():
                    cell_text += paragraph.text + " "
            row_data.append(cell_text.strip())
        table_data.append(row_data)
    return table_data

def get_docx_elements_in_order(document):
    """按顺序获取文档元素"""
    body = document._element.body
    elements = []
    for child in body.iterchildren():
        if child.tag.endswith('p'):
            for paragraph in document.paragraphs:
                if paragraph._element is child:
                    elements.append(('paragraph', paragraph))
                    break
        elif child.tag.endswith('tbl'):
            for table in document.tables:
                if table._element is child:
                    elements.append(('table', table))
                    break
    return elements

def parse_docx_to_json(docx_path):
    """解析docx文档为JSON结构"""
    doc = Document(docx_path)
    elements = get_docx_elements_in_order(doc)
    result = []
    heading_stack = [(-1, None, result)]
    current_content = []

    for elem_type, elem in elements:
        if elem_type == 'paragraph':
            level = get_paragraph_heading_level(elem)
            
            if level > 0:
                if heading_stack[-1][1] is not None:
                    heading_stack[-1][1]["content"] = current_content
                
                heading_node = {
                    "title": elem.text,
                    "content": []
                }
                current_content = []
                
                while heading_stack[-1][0] >= level:
                    heading_stack.pop()
                
                parent_level, parent_node, parent_container = heading_stack[-1]
                
                if isinstance(parent_container, list):
                    parent_container.append(heading_node)
                else:
                    if "children" not in parent_container:
                        parent_container["children"] = []
                    parent_container["children"].append(heading_node)
                
                heading_stack.append((level, heading_node, heading_node))
            elif elem.text.strip():
                current_content.append({"type": "text", "value": elem.text})
        
        elif elem_type == 'table':
            table_data = process_table(elem)
            current_content.append({"type": "table", "value": table_data})
    
    if heading_stack[-1][1] is not None:
        heading_stack[-1][1]["content"] = current_content
    
    return result

def transform_json(json_data):
    """转换JSON结构为漏洞报告格式"""
    result = {}
    issue_level_pattern = re.compile(r"【(低危|中危|高危|提示)】.*?漏洞数：(\d+)")

    for item in json_data:
        main_title = item["title"]
        result[main_title] = []
        if "children" not in item:
            continue
            
        for child in item["children"]:
            match = issue_level_pattern.search(child["title"])
            if match:
                issue_level = match.group(1)
                issue_count = match.group(2)

                if issue_level == "提示":
                    issue_level = "Notice"
                elif issue_level == "低危":
                    issue_level = "Low"
                elif issue_level == "中危":
                    issue_level = "Medium"
                elif issue_level == "高危":
                    issue_level = "High"
                    
                if issue_level not in ["Medium", "High"]:
                    continue

                issue = {
                    "issue_title": child["title"],
                    "issue_level": issue_level,
                    "issue_count": issue_count,
                    "issue_desc": "",
                    "fix_advice": "",
                    "code_sample": "",
                    "code_list": []
                }
                
                for section in child["children"]:
                    section_title = section["title"]
                    section_content = section["content"]
                    
                    if section_title == "漏洞描述":
                        for content_item in section_content:
                            if content_item["type"] == "text":
                                issue["issue_desc"] = content_item["value"]
                    
                    elif section_title == "修复建议":
                        for content_item in section_content:
                            if content_item["type"] == "text":
                                issue["fix_advice"] = content_item["value"]
                    
                    elif section_title == "代码示例":
                        for content_item in section_content:
                            if content_item["type"] == "text":
                                issue["code_sample"] = content_item["value"]
                    
                    elif re.match(r'^NO\.\d+\.\s代码位置$', section_title):
                        code_item = {}
                        
                        if section_content and section_content[0]["type"] == "text":
                            code_location_num = section_content[0]["value"]
                            splitor = code_location_num.split(":")
                            code_item["code_location"] = splitor[0]
                            code_item["code_line_num"] = splitor[1]

                        # 添加标准字段
                        code_item["status"] = "pending"
                        code_item["false_positive_probability"] = 0
                        code_item["false_positive_reason"] = ""
                        
                        if "children" in section:
                            for child_section in section["children"]:
                                for content_item in child_section["content"]:
                                    if content_item["type"] == "table":
                                        code_item["code_details"] = content_item["value"][-1][1]
                        
                        issue["code_list"].append(code_item)
                
                result[main_title].append(issue)
    
    return result

def convert_sast_docx_to_json(file_path: str, working_directory: str) -> str:
    """转换SAST docx文档为JSON"""
    # --- 修复点开始：对 file_path 做规范化与边界校验，防止路径遍历 ---
    wd_abs = os.path.abspath(working_directory)
    # 若为相对路径，先拼到工作目录；无论绝对或相对，统一做规范化
    candidate = os.path.abspath(os.path.join(wd_abs, file_path) if not os.path.isabs(file_path) else file_path)
    # 校验规范化后的路径是否仍在工作目录内（防止 ../../ 越界）
    try:
        if os.path.commonpath([wd_abs, candidate]) != wd_abs:
            return f"非法文件路径: {file_path}"
    except ValueError:
        # 在不同盘符/根路径等极端情况下的保护
        return f"非法文件路径: {file_path}"
    # 后续逻辑继续使用安全后的路径
    file_path = candidate
    # --- 修复点结束 ---

    # 如果文件不存在则返回
    if not os.path.exists(file_path):
        return f"文件不存在: {file_path}"

    try:
        # 使用多厂商解析器工厂解析文档
        parsed_result = sast_parser_factory.parse_docx(file_path)
        result = parsed_result.get("vulnerabilities", [])

        # 记录使用的解析器厂商
        vendor = parsed_result.get("vendor", "unknown")
        print(f"使用 {vendor} 解析器处理文档")

    except ValueError as e:
        # 没有找到合适的解析器，尝试使用默认逻辑
        print(f"警告: {str(e)}")
        print("尝试使用默认解析逻辑...")

        # 回退到原有解析逻辑
        all_docx_json = parse_docx_to_json(file_path)
        result = []
        for item in all_docx_json:
            if item["title"] in ["四、漏洞详情", "六、代码规范风险详情"]:
                result += transform_json([item])[item["title"]]
    except Exception as e:
        return f"解析文档时发生错误: {str(e)}"

    json_result = {"status": "success", "data": result}

    # 统一为所有厂商的解析结果添加标准字段
    for issue in json_result["data"]:
        for code_item in issue.get("code_list", []):
            # 添加缺失的标准字段
            if "status" not in code_item:
                code_item["status"] = "pending"
            if "false_positive_probability" not in code_item:
                code_item["false_positive_probability"] = 0
            if "false_positive_reason" not in code_item:
                code_item["false_positive_reason"] = ""

    output_dir = os.path.join(working_directory, ".scanissuefix")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for index, issue in enumerate(json_result["data"], start=1):
        issue_title = issue["issue_title"]
        issue_level = issue.get("issue_level", "未知")
        issue_count = issue.get("issue_count", "0")
        code_list = issue["code_list"]
        total_code_entries = len(code_list)

        # 转换风险级别为中文
        level_map = {"High": "高危", "Medium": "中危", "Low": "低危"}
        level_chinese = level_map.get(issue_level, issue_level)

        if total_code_entries > 5:
            chunk_size = 5
            chunks = [code_list[i:i + chunk_size] for i in range(0, total_code_entries, chunk_size)]

            for chunk_index, chunk in enumerate(chunks, start=1):
                issue_copy = issue.copy()
                issue_copy["code_list"] = chunk

                filename = f"{index}_{issue_title.replace(' ', '-').replace('【', '_').replace('】', '').replace('(', '_').replace(')', '').replace('：', '').replace(':', '').replace('/', '')}_{level_chinese}_漏洞数{issue_count}_{chunk_index}_new.json"
                file_path = os.path.join(output_dir, filename)

                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(issue_copy, f, ensure_ascii=False, indent=2)
        else:
            filename = f"{index}_{issue_title.replace(' ', '-').replace('【', '_').replace('】', '').replace('(', '_').replace(')', '').replace('：', '').replace(':', '').replace('/', '')}_{level_chinese}_漏洞数{issue_count}_new.json"
            file_path = os.path.join(output_dir, filename)
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(issue, f, ensure_ascii=False, indent=2)

    return f"所有漏洞已保存到{output_dir}目录"

def get_pending_vulnerability_json_files(working_directory: str) -> str:
    """获取待处理的漏洞JSON文件"""
    output_dir = os.path.join(working_directory, ".scanissuefix")
    if not os.path.exists(output_dir):
        return f"目录'{output_dir}'不存在"
    
    all_files = os.listdir(output_dir)
    new_json_files = [os.path.join(output_dir, file) for file in all_files if file.endswith("_new.json")]

    if not new_json_files:
        return f"目录'{output_dir}'中没有找到'_new.json'文件"
    
    return "\n".join(new_json_files)

def generate_csv_report(working_directory: str) -> str:
    """生成CSV报告"""
    output_dir = os.path.join(working_directory, ".scanissuefix")
    if not os.path.exists(output_dir):
        return f"目录'{output_dir}'不存在"
    
    all_files = os.listdir(output_dir)
    finished_json_files = [os.path.join(output_dir, file) for file in all_files if file.endswith("_finished.json")]

    if not finished_json_files:
        return f"目录'{output_dir}'中没有找到'_finished.json'文件"
    
    report_data = []
    for file_path in finished_json_files:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

            for code_item in data.get("code_list", []):
                report_data.append({
                    "漏洞类型": data["issue_title"],
                    "漏洞等级": data["issue_level"],
                    "代码位置": code_item["code_location"],
                    "代码行号": code_item["code_line_num"],
                    "代码详情": code_item["code_details"],
                    "修复状态": code_item.get("status", "missed"),
                    "误报概率": code_item.get("false_positive_probability", ""),
                    "误报原因澄清": code_item.get("false_positive_reason", "")
                })

    csv_output_path = os.path.join(output_dir, "sast_fix_report.csv")

    with open(csv_output_path, 'w', newline='', encoding='utf-8') as csv_file:
        fieldnames = [
            "漏洞类型", 
            "漏洞等级", 
            "代码位置", 
            "代码行号", 
            "代码详情", 
            "修复状态", 
            "误报概率", 
            "误报原因澄清"
        ]
        writer = csv.DictWriter(csv_file, fieldnames=fieldnames, quoting=csv.QUOTE_MINIMAL)
        writer.writeheader()
        writer.writerows(report_data)

    return f"CSV报告生成成功: {csv_output_path}"

async def serve(working_directory: Path | None) -> None:
    logger = logging.getLogger(__name__)

    if working_directory is not None and not working_directory.exists():
        logger.error(f"Working directory {working_directory} does not exist")
        return

    server = Server("sast-fixer-mcp")

    @server.list_tools()
    async def list_tools() -> List[Tool]:
        return [
            Tool(
                name=SastTools.CONVERT_DOCX_TO_JSON,
                description="将SAST报告的docx文档转换为JSON格式，使用服务器配置的工作目录",
                inputSchema=ConvertSastDocxToJson.model_json_schema(),
            ),
            Tool(
                name=SastTools.GET_PENDING_FILES,
                description="获取工作目录下.scanissuefix中所有待处理的漏洞JSON文件(_new.json)",
                inputSchema=GetPendingVulnerabilityJsonFiles.model_json_schema(),
            ),
            Tool(
                name=SastTools.GENERATE_CSV_REPORT,
                description="从工作目录下.scanissuefix中已完成的漏洞JSON文件(_finished.json)生成CSV报告",
                inputSchema=GenerateCsvReport.model_json_schema(),
            )
        ]

    @server.call_tool()
    async def call_tool(name: str, arguments: dict) -> List[TextContent]:
        try:            
            # 优先使用参数中的工作目录，然后是命令行传入的目录，最后使用当前目录
            work_dir = (
                arguments.get("working_directory") or 
                (str(working_directory) if working_directory else None) or 
                os.getcwd()
            )
            # # 使用命令行指定的工作目录，默认为当前目录
            # work_dir = str(working_directory) if working_directory else os.getcwd()
            
            match name:
                case SastTools.CONVERT_DOCX_TO_JSON:
                    result = convert_sast_docx_to_json(arguments["file_path"], work_dir)
                    return [TextContent(type="text", text=result)]

                case SastTools.GET_PENDING_FILES:
                    result = get_pending_vulnerability_json_files(work_dir)
                    return [TextContent(type="text", text=result)]

                case SastTools.GENERATE_CSV_REPORT:
                    result = generate_csv_report(work_dir)
                    return [TextContent(type="text", text=result)]

                case _:
                    raise ValueError(f"Unknown tool: {name}")
                    
        except Exception as e:
            error_msg = f"工具执行失败: {str(e)}"
            logger.error(error_msg)
            return [TextContent(type="text", text=error_msg)]

    options = server.create_initialization_options()
    async with stdio_server() as (read_stream, write_stream):
        await server.run(read_stream, write_stream, options, raise_exceptions=True)

if __name__ == "__main__":
    working_directory = Path(os.getcwd())
    temp_test_file = "mcp-servers/servers/sast_fixer_mcp/tests/sast_report_test.docx"
    result = convert_sast_docx_to_json(temp_test_file, working_directory)
    print(result)