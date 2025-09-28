#!/usr/bin/env python3
"""
Output Formatters - Multiple export format support

This module provides formatters for different output types:
- Plain text with timestamps
- Markdown with speaker labels
- Microsoft Word document
"""

import json
import logging
from abc import ABC, abstractmethod
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class TranscriptionSegment:
    """Represents a single segment of transcription."""

    start_time: float
    end_time: float
    text: str
    speaker: Optional[str] = None
    confidence: Optional[float] = None


@dataclass
class TranscriptionResult:
    """Complete transcription result."""

    segments: List[TranscriptionSegment]
    language: Optional[str] = None
    processing_time: Optional[float] = None
    metadata: Optional[Dict[str, Any]] = None


class BaseFormatter(ABC):
    """Abstract base class for output formatters."""

    def __init__(self, output_path: Path):
        """Initialize formatter.

        Args:
            output_path: Path where to save the formatted output
        """
        self.output_path = output_path
        self.logger = logging.getLogger(
            f"offline_stenographer.formatters.{self.__class__.__name__}"
        )

    @abstractmethod
    def format_transcription(self, result: TranscriptionResult) -> bool:
        """Format and save transcription result.

        Args:
            result: Transcription result to format

        Returns:
            True if successful, False otherwise
        """
        pass

    def _format_timestamp(self, seconds: float) -> str:
        """Format timestamp as HH:MM:SS.

        Args:
            seconds: Time in seconds

        Returns:
            Formatted timestamp string
        """
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        seconds = int(seconds % 60)

        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
        else:
            return f"{minutes:02d}:{seconds:02d}"

    def _consolidate_segments(
        self, segments: List[TranscriptionSegment]
    ) -> Dict[str, List[TranscriptionSegment]]:
        """Consolidate consecutive segments from the same speaker.

        Args:
            segments: List of transcription segments

        Returns:
            Dictionary mapping speaker names to their consolidated segments
        """
        consolidated = {}
        current_speaker_segments = []
        current_speaker = None

        for segment in segments:
            if segment.speaker == current_speaker:
                # Same speaker, add to current group
                current_speaker_segments.append(segment)
            else:
                # Speaker changed, save previous group
                if current_speaker and current_speaker_segments:
                    consolidated[current_speaker] = current_speaker_segments

                # Start new group
                current_speaker = segment.speaker
                current_speaker_segments = [segment]

        # Don't forget the last group
        if current_speaker and current_speaker_segments:
            consolidated[current_speaker] = current_speaker_segments

        return consolidated


class TextFormatter(BaseFormatter):
    """Plain text formatter with timestamps and dialog consolidation."""

    def format_transcription(self, result: TranscriptionResult) -> bool:
        """Format as plain text with consolidated dialog segments.

        Args:
            result: Transcription result to format

        Returns:
            True if successful
        """
        try:
            with open(self.output_path, "w", encoding="utf-8") as f:
                f.write("TRANSCRIPTION (DIALOG FORMAT)\n")
                f.write("=" * 50 + "\n\n")

                # Add metadata if available
                if result.metadata:
                    f.write("METADATA\n")
                    f.write("-" * 20 + "\n")
                    for key, value in result.metadata.items():
                        f.write(f"{key}: {value}\n")
                    f.write("\n")

                # Add consolidated dialog segments
                f.write("DIALOG\n")
                f.write("-" * 20 + "\n\n")

                # Consolidate consecutive segments from the same speaker
                consolidated_segments = self._consolidate_segments(result.segments)

                for speaker, segments in consolidated_segments.items():
                    # Speaker header
                    f.write(f"{speaker}:\n")

                    # Combine all segments for this speaker
                    combined_text = " ".join(segment.text for segment in segments)

                    # Add timestamp range if multiple segments
                    if len(segments) > 1:
                        start_time = segments[0].start_time
                        end_time = segments[-1].end_time
                        timestamp = f"[{self._format_timestamp(start_time)} - {self._format_timestamp(end_time)}]"
                    else:
                        timestamp = (
                            f"[{self._format_timestamp(segments[0].start_time)}]"
                        )

                    f.write(f"{timestamp} {combined_text}\n\n")

                self.logger.info(f"Text transcription saved to {self.output_path}")
                return True

        except Exception as e:
            self.logger.error(f"Error saving text format: {e}")
            return False


class MarkdownFormatter(BaseFormatter):
    """Markdown formatter with enhanced formatting and dialog consolidation."""

    def format_transcription(self, result: TranscriptionResult) -> bool:
        """Format as Markdown with consolidated dialog segments.

        Args:
            result: Transcription result to format

        Returns:
            True if successful
        """
        try:
            with open(self.output_path, "w", encoding="utf-8") as f:
                # Header
                f.write("# Transcription (Dialog Format)\n\n")
                f.write(
                    f"*Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n"
                )

                # Metadata section
                if result.metadata:
                    f.write("## Metadata\n\n")
                    f.write("| Property | Value |\n")
                    f.write("|----------|-------|\n")
                    for key, value in result.metadata.items():
                        f.write(f"| {key} | {value} |\n")
                    f.write("\n")

                # Content section
                f.write("## Dialog\n\n")

                # Consolidate consecutive segments from the same speaker
                consolidated_segments = self._consolidate_segments(result.segments)

                for speaker, segments in consolidated_segments.items():
                    # Speaker header
                    f.write(f"### {speaker}\n\n")

                    # Combine all segments for this speaker
                    combined_text = " ".join(segment.text for segment in segments)

                    # Add timestamp range if multiple segments
                    if len(segments) > 1:
                        start_time = segments[0].start_time
                        end_time = segments[-1].end_time
                        timestamp = f"**[{self._format_timestamp(start_time)} - {self._format_timestamp(end_time)}]**"
                    else:
                        timestamp = (
                            f"**[{self._format_timestamp(segments[0].start_time)}]**"
                        )

                    # Segment content
                    f.write(f"{timestamp} {combined_text}\n\n")

                self.logger.info(f"Markdown transcription saved to {self.output_path}")
                return True

        except Exception as e:
            self.logger.error(f"Error saving markdown format: {e}")
            return False


class DocxFormatter(BaseFormatter):
    """Microsoft Word document formatter with dialog consolidation."""

    def format_transcription(self, result: TranscriptionResult) -> bool:
        """Format as Microsoft Word document with consolidated dialog segments.

        Args:
            result: Transcription result to format

        Returns:
            True if successful
        """
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx not available for DOCX formatting")
            return False

        try:
            doc = Document()

            # Title
            title = doc.add_heading("Transcription (Dialog Format)", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata
            if result.metadata:
                doc.add_heading("Metadata", level=1)
                table = doc.add_table(rows=0, cols=2)

                for key, value in result.metadata.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(key)
                    row_cells[1].text = str(value)

                doc.add_paragraph()  # Add spacing

            # Content
            doc.add_heading("Dialog", level=1)

            # Consolidate consecutive segments from the same speaker
            consolidated_segments = self._consolidate_segments(result.segments)

            for speaker, segments in consolidated_segments.items():
                # Speaker header
                speaker_heading = doc.add_heading(speaker, level=2)
                speaker_heading.style = "Heading 2"

                # Combine all segments for this speaker
                combined_text = " ".join(segment.text for segment in segments)

                # Add timestamp range if multiple segments
                if len(segments) > 1:
                    start_time = segments[0].start_time
                    end_time = segments[-1].end_time
                    timestamp = f"[{self._format_timestamp(start_time)} - {self._format_timestamp(end_time)}]"
                else:
                    timestamp = f"[{self._format_timestamp(segments[0].start_time)}]"

                # Segment content
                para = doc.add_paragraph()
                timestamp_run = para.add_run(f"{timestamp} ")
                timestamp_run.bold = True
                para.add_run(combined_text)

                doc.add_paragraph()  # Add spacing between speakers

            # Save document
            doc.save(str(self.output_path))
            self.logger.info(f"DOCX transcription saved to {self.output_path}")
            return True

        except Exception as e:
            self.logger.error(f"Error saving DOCX format: {e}")
            return False


class FormatterFactory:
    """Factory for creating appropriate formatters."""

    FORMATTERS = {"txt": TextFormatter, "md": MarkdownFormatter, "docx": DocxFormatter}

    @classmethod
    def create_formatter(
        cls, format_type: str, output_path: Path
    ) -> Optional[BaseFormatter]:
        """Create formatter for specified format.

        Args:
            format_type: Type of formatter ('txt', 'md', 'docx')
            output_path: Path where to save output

        Returns:
            Formatter instance or None if format not supported
        """
        formatter_class = cls.FORMATTERS.get(format_type.lower())
        if formatter_class is None:
            logging.error(f"Unsupported format: {format_type}")
            return None

        return formatter_class(output_path)

    @classmethod
    def get_supported_formats(cls) -> List[str]:
        """Get list of supported output formats.

        Returns:
            List of supported format strings
        """
        return list(cls.FORMATTERS.keys())


def format_transcription_output(
    result: TranscriptionResult, output_format: str, output_path: Path
) -> bool:
    """Format transcription result to specified format.

    Args:
        result: Transcription result to format
        output_format: Output format ('txt', 'md', 'docx')
        output_path: Path where to save formatted output

    Returns:
        True if successful, False otherwise
    """
    formatter = FormatterFactory.create_formatter(output_format, output_path)
    if formatter is None:
        return False

    return formatter.format_transcription(result)
