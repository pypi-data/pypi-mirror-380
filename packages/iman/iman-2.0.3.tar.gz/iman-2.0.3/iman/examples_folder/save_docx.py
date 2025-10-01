# python-docx==0.8.11
# lxml


from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH 


document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True
p.alignment=WD_ALIGN_PARAGRAPH.CENTER  # for left, 1 for center, 2 right, 3 justify ....

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')



doc = docx.Document()
style = doc.styles.add_style('rtl', WD_STYLE_TYPE.PARAGRAPH)
# style.font.rtl = True
style.font.name = 'Nazanin'
style.font.size = Pt(14)

p = doc.add_paragraph(style='rtl')
p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
p.add_run('بنام خدا')

doc.save('iman.docx')