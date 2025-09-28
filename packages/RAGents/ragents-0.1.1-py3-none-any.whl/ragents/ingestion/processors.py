"""Document processors for different file formats."""

import asyncio
import base64
import io
import json
from abc import ABC, abstractmethod
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from pydantic import BaseModel

from ..rag.types import Document
from .config import IngestionConfig


class DocumentProcessor(ABC):
    """Abstract base class for document processors."""

    @abstractmethod
    async def process(self, file_path: Path, config: IngestionConfig) -> Document:
        """Process a document and return a Document object."""
        pass

    @abstractmethod
    def can_process(self, file_path: Path) -> bool:
        """Check if this processor can handle the given file."""
        pass

    async def create_chunks(self, document: Document, config: IngestionConfig) -> List[Dict[str, Any]]:
        """Create chunks from the document."""
        content = document.content
        chunk_size = config.chunk_size
        chunk_overlap = config.chunk_overlap

        chunks = []
        start = 0

        while start < len(content):
            end = start + chunk_size
            chunk_content = content[start:end]

            if chunk_content.strip():  # Only create non-empty chunks
                chunks.append({
                    "content": chunk_content,
                    "metadata": {
                        **document.metadata,
                        "chunk_index": len(chunks),
                        "start_char": start,
                        "end_char": end,
                        "chunk_size": len(chunk_content)
                    }
                })

            start = end - chunk_overlap

        return chunks


class TextProcessor(DocumentProcessor):
    """Processor for plain text files."""

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in [".txt", ".md", ".rtf"]

    async def process(self, file_path: Path, config: IngestionConfig) -> Document:
        """Process a text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                content = f.read()

        return Document(
            id=str(file_path),
            content=content,
            source=str(file_path),
            metadata={
                "file_type": "text",
                "file_size": file_path.stat().st_size,
                "created_at": datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
                "modified_at": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
                "encoding": "utf-8",
                "line_count": content.count('\n') + 1,
                "word_count": len(content.split()),
                "char_count": len(content)
            }
        )


class PDFProcessor(DocumentProcessor):
    """Processor for PDF files."""

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".pdf"

    async def process(self, file_path: Path, config: IngestionConfig) -> Document:
        """Process a PDF file."""
        try:
            import pymupdf as fitz
        except ImportError:
            try:
                import fitz
            except ImportError:
                raise ImportError("PyMuPDF is required for PDF processing. Install with: pip install pymupdf")

        doc = fitz.open(file_path)
        content_parts = []
        images = []
        tables = []

        metadata = {
            "file_type": "pdf",
            "file_size": file_path.stat().st_size,
            "page_count": len(doc),
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
            "subject": doc.metadata.get("subject", ""),
            "created_at": doc.metadata.get("creationDate", ""),
            "modified_at": doc.metadata.get("modDate", ""),
        }

        for page_num, page in enumerate(doc):
            # Extract text
            text = page.get_text()
            content_parts.append(f"--- Page {page_num + 1} ---\n{text}")

            # Extract images if enabled
            if config.extract_images:
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        if pix.n - pix.alpha < 4:  # GRAY or RGB
                            img_data = pix.tobytes("png")
                            images.append({
                                "page": page_num + 1,
                                "index": img_index,
                                "data": base64.b64encode(img_data).decode(),
                                "format": "png"
                            })
                        pix = None
                    except Exception as e:
                        print(f"Error extracting image: {e}")

            # Extract tables if enabled
            if config.extract_tables:
                try:
                    tables_on_page = page.find_tables()
                    for table_index, table in enumerate(tables_on_page):
                        table_data = table.extract()
                        tables.append({
                            "page": page_num + 1,
                            "index": table_index,
                            "data": table_data
                        })
                except Exception as e:
                    print(f"Error extracting table: {e}")

        doc.close()

        content = "\n\n".join(content_parts)

        # Add extracted data to metadata
        if images:
            metadata["images"] = images
        if tables:
            metadata["tables"] = tables

        metadata.update({
            "word_count": len(content.split()),
            "char_count": len(content),
            "images_count": len(images),
            "tables_count": len(tables)
        })

        return Document(
            id=str(file_path),
            content=content,
            source=str(file_path),
            metadata=metadata
        )


class CSVProcessor(DocumentProcessor):
    """Processor for CSV files."""

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".csv"

    async def process(self, file_path: Path, config: IngestionConfig) -> Document:
        """Process a CSV file."""
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("pandas is required for CSV processing. Install with: pip install pandas")

        try:
            df = pd.read_csv(file_path)
        except Exception as e:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except:
                    continue
            else:
                raise e

        # Create content representation
        content_parts = [
            f"CSV Dataset: {file_path.name}",
            f"Columns: {', '.join(df.columns.tolist())}",
            f"Rows: {len(df)}",
            "",
            "Column Information:",
        ]

        # Add column descriptions
        for col in df.columns:
            dtype = str(df[col].dtype)
            null_count = df[col].isnull().sum()
            unique_count = df[col].nunique()

            content_parts.append(
                f"- {col}: {dtype}, {null_count} nulls, {unique_count} unique values"
            )

        # Add sample data
        content_parts.extend([
            "",
            "Sample Data (first 5 rows):",
            df.head().to_string(index=False)
        ])

        # Add summary statistics for numeric columns
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            content_parts.extend([
                "",
                "Numeric Column Statistics:",
                df[numeric_cols].describe().to_string()
            ])

        content = "\n".join(content_parts)

        metadata = {
            "file_type": "csv",
            "file_size": file_path.stat().st_size,
            "row_count": len(df),
            "column_count": len(df.columns),
            "columns": df.columns.tolist(),
            "column_types": {col: str(df[col].dtype) for col in df.columns},
            "null_counts": df.isnull().sum().to_dict(),
            "memory_usage": df.memory_usage(deep=True).sum(),
            "created_at": datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
        }

        return Document(
            id=str(file_path),
            content=content,
            source=str(file_path),
            metadata=metadata
        )


class ParquetProcessor(DocumentProcessor):
    """Processor for Parquet files."""

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".parquet"

    async def process(self, file_path: Path, config: IngestionConfig) -> Document:
        """Process a Parquet file."""
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("pandas is required for Parquet processing. Install with: pip install pandas")

        try:
            import pyarrow.parquet as pq
        except ImportError:
            raise ImportError("pyarrow is required for Parquet processing. Install with: pip install pyarrow")

        df = pd.read_parquet(file_path)

        # Get parquet metadata
        parquet_file = pq.ParquetFile(file_path)
        parquet_metadata = parquet_file.metadata

        # Create content representation
        content_parts = [
            f"Parquet Dataset: {file_path.name}",
            f"Columns: {', '.join(df.columns.tolist())}",
            f"Rows: {len(df)}",
            f"Row Groups: {parquet_metadata.num_row_groups}",
            "",
            "Schema Information:",
        ]

        # Add schema details
        schema = parquet_file.schema_arrow
        for i, field in enumerate(schema):
            content_parts.append(f"- {field.name}: {field.type}")

        # Add sample data
        content_parts.extend([
            "",
            "Sample Data (first 5 rows):",
            df.head().to_string(index=False)
        ])

        # Add summary statistics
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            content_parts.extend([
                "",
                "Numeric Column Statistics:",
                df[numeric_cols].describe().to_string()
            ])

        content = "\n".join(content_parts)

        metadata = {
            "file_type": "parquet",
            "file_size": file_path.stat().st_size,
            "row_count": len(df),
            "column_count": len(df.columns),
            "row_groups": parquet_metadata.num_row_groups,
            "columns": df.columns.tolist(),
            "schema": {field.name: str(field.type) for field in schema},
            "parquet_version": parquet_metadata.format_version,
            "created_by": parquet_metadata.created_by,
            "memory_usage": df.memory_usage(deep=True).sum(),
            "created_at": datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
        }

        return Document(
            id=str(file_path),
            content=content,
            source=str(file_path),
            metadata=metadata
        )


class ImageProcessor(DocumentProcessor):
    """Processor for image files."""

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"]

    async def process(self, file_path: Path, config: IngestionConfig) -> Document:
        """Process an image file."""
        try:
            from PIL import Image
        except ImportError:
            raise ImportError("Pillow is required for image processing. Install with: pip install Pillow")

        # Load image
        with Image.open(file_path) as img:
            # Get basic image information
            width, height = img.size
            mode = img.mode
            format_name = img.format

            # Convert to RGB if necessary for analysis
            if mode != 'RGB':
                img_rgb = img.convert('RGB')
            else:
                img_rgb = img

            # Extract image data
            img_bytes = io.BytesIO()
            img_rgb.save(img_bytes, format='PNG')
            img_base64 = base64.b64encode(img_bytes.getvalue()).decode()

        # Create textual description of the image
        content_parts = [
            f"Image: {file_path.name}",
            f"Dimensions: {width}x{height} pixels",
            f"Format: {format_name}",
            f"Mode: {mode}",
            "",
            "This is an image file. Use image analysis tools to extract detailed content information.",
        ]

        # TODO: Add OCR if text is detected
        if config.enable_ocr:
            try:
                import pytesseract
                ocr_text = pytesseract.image_to_string(img_rgb, lang=config.ocr_language)
                if ocr_text.strip():
                    content_parts.extend([
                        "",
                        "Extracted Text (OCR):",
                        ocr_text
                    ])
            except ImportError:
                print("pytesseract not available for OCR")
            except Exception as e:
                print(f"OCR extraction failed: {e}")

        content = "\n".join(content_parts)

        metadata = {
            "file_type": "image",
            "file_size": file_path.stat().st_size,
            "width": width,
            "height": height,
            "format": format_name,
            "mode": mode,
            "aspect_ratio": width / height,
            "total_pixels": width * height,
            "image_data": img_base64,
            "created_at": datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
        }

        return Document(
            id=str(file_path),
            content=content,
            source=str(file_path),
            metadata=metadata
        )


class JSONProcessor(DocumentProcessor):
    """Processor for JSON files."""

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".json"

    async def process(self, file_path: Path, config: IngestionConfig) -> Document:
        """Process a JSON file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Create structured content representation
        content_parts = [
            f"JSON Document: {file_path.name}",
            "",
            "Structure Overview:",
            self._analyze_json_structure(data),
            "",
            "Full Content:",
            json.dumps(data, indent=2, ensure_ascii=False)
        ]

        content = "\n".join(content_parts)

        metadata = {
            "file_type": "json",
            "file_size": file_path.stat().st_size,
            "json_structure": self._get_json_schema(data),
            "key_count": self._count_keys(data),
            "max_depth": self._get_max_depth(data),
            "created_at": datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
        }

        return Document(
            id=str(file_path),
            content=content,
            source=str(file_path),
            metadata=metadata
        )

    def _analyze_json_structure(self, data: Any, prefix: str = "") -> str:
        """Analyze JSON structure and return a description."""
        if isinstance(data, dict):
            lines = [f"{prefix}Object with {len(data)} keys:"]
            for key in list(data.keys())[:10]:  # Show first 10 keys
                value_type = type(data[key]).__name__
                lines.append(f"{prefix}  - {key}: {value_type}")
            if len(data) > 10:
                lines.append(f"{prefix}  ... and {len(data) - 10} more keys")
            return "\n".join(lines)
        elif isinstance(data, list):
            return f"{prefix}Array with {len(data)} items"
        else:
            return f"{prefix}{type(data).__name__}: {str(data)[:100]}"

    def _get_json_schema(self, data: Any) -> Dict[str, Any]:
        """Get simplified schema of JSON data."""
        if isinstance(data, dict):
            return {key: self._get_json_schema(value) for key, value in data.items()}
        elif isinstance(data, list):
            if data:
                return [self._get_json_schema(data[0])]
            else:
                return []
        else:
            return type(data).__name__

    def _count_keys(self, data: Any) -> int:
        """Count total number of keys in nested structure."""
        if isinstance(data, dict):
            return len(data) + sum(self._count_keys(v) for v in data.values())
        elif isinstance(data, list):
            return sum(self._count_keys(item) for item in data)
        else:
            return 0

    def _get_max_depth(self, data: Any, current_depth: int = 0) -> int:
        """Get maximum nesting depth."""
        if isinstance(data, dict):
            if not data:
                return current_depth
            return max(self._get_max_depth(v, current_depth + 1) for v in data.values())
        elif isinstance(data, list):
            if not data:
                return current_depth
            return max(self._get_max_depth(item, current_depth + 1) for item in data)
        else:
            return current_depth


class DocxProcessor(DocumentProcessor):
    """Processor for Microsoft Word documents."""

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in [".docx", ".doc"]

    async def process(self, file_path: Path, config: IngestionConfig) -> Document:
        """Process a Word document."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")

        doc = DocxDocument(file_path)

        # Extract text content
        content_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content_parts.append(para.text)

        # Extract tables if enabled
        tables_data = []
        if config.extract_tables:
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append({
                    "index": table_idx,
                    "data": table_data
                })

                # Add table to content
                content_parts.append(f"\n--- Table {table_idx + 1} ---")
                for row in table_data:
                    content_parts.append(" | ".join(row))

        content = "\n\n".join(content_parts)

        # Extract document properties
        props = doc.core_properties
        metadata = {
            "file_type": "docx",
            "file_size": file_path.stat().st_size,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "word_count": len(content.split()),
            "char_count": len(content),
            "title": props.title or "",
            "author": props.author or "",
            "subject": props.subject or "",
            "created": props.created.isoformat() if props.created else "",
            "modified": props.modified.isoformat() if props.modified else "",
            "tables": tables_data if tables_data else [],
            "created_at": datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
        }

        return Document(
            id=str(file_path),
            content=content,
            source=str(file_path),
            metadata=metadata
        )


# Registry of processors
_PROCESSORS = [
    TextProcessor(),
    PDFProcessor(),
    CSVProcessor(),
    ParquetProcessor(),
    ImageProcessor(),
    JSONProcessor(),
    DocxProcessor(),
]


def get_processor_for_file(file_path: Path) -> Optional[DocumentProcessor]:
    """Get the appropriate processor for a file."""
    for processor in _PROCESSORS:
        if processor.can_process(file_path):
            return processor
    return None


def get_supported_formats() -> List[str]:
    """Get list of all supported file formats."""
    formats = set()
    for processor in _PROCESSORS:
        if hasattr(processor, 'can_process'):
            # This is a simplified approach - in reality you'd want to
            # have each processor declare its supported formats
            pass

    return [
        ".txt", ".md", ".rtf",  # Text
        ".pdf",  # PDF
        ".csv",  # CSV
        ".parquet",  # Parquet
        ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp",  # Images
        ".json",  # JSON
        ".docx", ".doc",  # Word documents
    ]


def register_processor(processor: DocumentProcessor):
    """Register a new document processor."""
    _PROCESSORS.append(processor)


class MultiModalProcessor:
    """Unified processor for handling multiple file types with multimodal capabilities."""

    def __init__(self, config: IngestionConfig):
        """Initialize the multimodal processor."""
        self.config = config
        self.processors = {
            processor.__class__.__name__: processor
            for processor in _PROCESSORS
        }

    async def process_file(self, file_path: Path) -> Optional[Document]:
        """Process a file using the appropriate processor."""
        processor = get_processor_for_file(file_path)
        if processor:
            return await processor.process(file_path, self.config)
        return None

    async def process_files(self, file_paths: List[Path]) -> List[Document]:
        """Process multiple files concurrently."""
        tasks = [self.process_file(path) for path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        documents = []
        for result in results:
            if isinstance(result, Document):
                documents.append(result)
            elif isinstance(result, Exception):
                print(f"Error processing file: {result}")

        return documents

    def get_supported_formats(self) -> List[str]:
        """Get all supported file formats."""
        return get_supported_formats()