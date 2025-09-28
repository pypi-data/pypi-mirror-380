"""
Tests for core functionality of docdocgo package.
"""

import os
import tempfile
from pathlib import Path

import pandas as pd
import pytest
from docx import Document

from docdocgo.core import create_document_with_table, docdocgo, insert_table


@pytest.fixture
def sample_dataframe():
    """Create a sample DataFrame for testing."""
    return pd.DataFrame(
        {
            "Name": ["Alice", "Bob", "Charlie"],
            "Age": [25, 30, 35],
            "Score": [85.5, 92.0, 78.3],
        }
    )


@pytest.fixture
def temp_dir():
    """Create a temporary directory for test files."""
    with tempfile.TemporaryDirectory() as tmp_dir:
        yield Path(tmp_dir)


@pytest.fixture
def sample_docx_with_marker(temp_dir):
    """Create a sample Word document with a marker."""
    doc = Document()
    doc.add_paragraph("This is a test document.")
    doc.add_paragraph("Here is the data: {{TEST_MARKER}}")
    doc.add_paragraph("End of document.")

    doc_path = temp_dir / "test_with_marker.docx"
    doc.save(doc_path)
    return doc_path


class TestCreateDocumentWithTable:
    """Test create_document_with_table function."""

    def test_create_new_document(self, sample_dataframe, temp_dir):
        """Test creating a new document with a table."""
        output_path = temp_dir / "new_document.docx"

        create_document_with_table(sample_dataframe, output_path)

        assert output_path.exists()

        # Verify the document content
        doc = Document(output_path)
        tables = doc.tables
        assert len(tables) == 1

        table = tables[0]
        assert len(table.rows) == 4  # 1 header + 3 data rows
        assert len(table.columns) == 3  # Name, Age, Score

    def test_empty_dataframe_raises_error(self, temp_dir):
        """Test that empty DataFrame raises ValueError."""
        empty_df = pd.DataFrame()
        output_path = temp_dir / "empty_test.docx"

        with pytest.raises(ValueError, match="DataFrame cannot be empty"):
            create_document_with_table(empty_df, output_path)

    def test_invalid_dataframe_raises_error(self, temp_dir):
        """Test that non-DataFrame input raises TypeError."""
        output_path = temp_dir / "invalid_test.docx"

        with pytest.raises(TypeError, match="Expected pandas DataFrame"):
            create_document_with_table("not a dataframe", output_path)


class TestInsertTable:
    """Test insert_table function."""

    def test_insert_table_with_marker(self, sample_dataframe, sample_docx_with_marker):
        """Test inserting table by replacing marker."""
        original_size = os.path.getsize(sample_docx_with_marker)

        insert_table(sample_dataframe, sample_docx_with_marker, "{{TEST_MARKER}}")

        # Document should be larger after adding table
        new_size = os.path.getsize(sample_docx_with_marker)
        assert new_size > original_size

        # Verify the document content
        doc = Document(sample_docx_with_marker)
        tables = doc.tables
        assert len(tables) == 1

        table = tables[0]
        assert len(table.rows) == 4  # 1 header + 3 data rows
        assert len(table.columns) == 3

    def test_marker_not_found_raises_error(
        self, sample_dataframe, sample_docx_with_marker
    ):
        """Test that missing marker raises ValueError."""
        with pytest.raises(ValueError, match="Marker '{{MISSING_MARKER}}' not found"):
            insert_table(
                sample_dataframe, sample_docx_with_marker, "{{MISSING_MARKER}}"
            )

    def test_file_not_found_raises_error(self, sample_dataframe, temp_dir):
        """Test that missing file raises FileNotFoundError."""
        nonexistent_path = temp_dir / "nonexistent.docx"

        with pytest.raises(FileNotFoundError):
            insert_table(sample_dataframe, nonexistent_path, "{{TEST_MARKER}}")

    def test_insert_table_with_custom_output(
        self, sample_dataframe, sample_docx_with_marker, temp_dir
    ):
        """Test inserting table with custom output path (preserves original)."""
        output_path = temp_dir / "custom_output.docx"

        # Original should remain unchanged
        original_content = sample_docx_with_marker.read_bytes()

        insert_table(
            sample_dataframe, sample_docx_with_marker, "{{TEST_MARKER}}", output_path
        )

        # Original file should be unchanged
        assert sample_docx_with_marker.read_bytes() == original_content

        # New file should exist and contain the table
        assert output_path.exists()
        doc = Document(output_path)
        assert len(doc.tables) == 1

    def test_marker_edge_cases(self, sample_dataframe, temp_dir):
        """Test edge cases for marker replacement."""
        # Create document with no markers
        doc = Document()
        doc.add_paragraph("This document has no markers.")
        doc_path = temp_dir / "no_markers.docx"
        doc.save(doc_path)

        # Should raise ValueError with clear message
        with pytest.raises(
            ValueError, match="Marker '{{MISSING_TAG}}' not found in document"
        ):
            insert_table(sample_dataframe, doc_path, "{{MISSING_TAG}}")

    def test_multiple_markers_in_document(self, sample_dataframe, temp_dir):
        """Test document with multiple markers (only first should be replaced)."""
        doc = Document()
        doc.add_paragraph("First marker: {{TEST_MARKER}}")
        doc.add_paragraph("Second marker: {{TEST_MARKER}}")
        doc_path = temp_dir / "multiple_markers.docx"
        doc.save(doc_path)

        insert_table(sample_dataframe, doc_path, "{{TEST_MARKER}}")

        # Verify only first marker was replaced
        doc = Document(doc_path)
        assert len(doc.tables) == 1

        # Check that second marker still exists
        paragraphs_text = [p.text for p in doc.paragraphs]
        remaining_markers = [p for p in paragraphs_text if "{{TEST_MARKER}}" in p]
        assert len(remaining_markers) == 1

    def test_case_sensitive_marker_search(self, sample_dataframe, temp_dir):
        """Test that marker search is case-sensitive."""
        doc = Document()
        doc.add_paragraph("Document contains: {{test_marker}}")  # lowercase
        doc_path = temp_dir / "case_sensitive.docx"
        doc.save(doc_path)

        # Should not find uppercase marker
        with pytest.raises(
            ValueError, match="Marker '{{TEST_MARKER}}' not found in document"
        ):
            insert_table(sample_dataframe, doc_path, "{{TEST_MARKER}}")

    def test_partial_marker_not_replaced(self, sample_dataframe, temp_dir):
        """Test that partial markers are not replaced."""
        doc = Document()
        doc.add_paragraph("Partial marker: {{TEST_MARKE")  # Missing closing
        doc.add_paragraph("Another partial: EST_MARKER}}")  # Missing opening
        doc_path = temp_dir / "partial_markers.docx"
        doc.save(doc_path)

        with pytest.raises(
            ValueError, match="Marker '{{TEST_MARKER}}' not found in document"
        ):
            insert_table(sample_dataframe, doc_path, "{{TEST_MARKER}}")


class TestDocdocgo:
    """Test main docdocgo function."""

    def test_docdocgo_with_marker(self, sample_dataframe, sample_docx_with_marker):
        """Test docdocgo with marker (insert mode)."""
        docdocgo(sample_dataframe, sample_docx_with_marker, "{{TEST_MARKER}}")

        # Verify the document has a table
        doc = Document(sample_docx_with_marker)
        assert len(doc.tables) == 1

    def test_docdocgo_without_marker(self, sample_dataframe, temp_dir):
        """Test docdocgo without marker (create mode)."""
        output_path = temp_dir / "docdocgo_test.docx"

        docdocgo(sample_dataframe, output_path)

        assert output_path.exists()

        # Verify the document has a table
        doc = Document(output_path)
        assert len(doc.tables) == 1

    def test_dataframe_with_missing_values(self, temp_dir):
        """Test handling DataFrame with NaN values."""
        df_with_nan = pd.DataFrame(
            {"A": [1, 2, None], "B": [4.5, None, 6.7], "C": ["x", "y", "z"]}
        )

        output_path = temp_dir / "nan_test.docx"
        docdocgo(df_with_nan, output_path)

        assert output_path.exists()

        # Verify document was created successfully
        doc = Document(output_path)
        assert len(doc.tables) == 1

    def test_docdocgo_with_custom_output_path(
        self, sample_dataframe, sample_docx_with_marker, temp_dir
    ):
        """Test docdocgo with custom output path."""
        output_path = temp_dir / "custom_docdocgo_output.docx"

        # Original should remain unchanged
        original_content = sample_docx_with_marker.read_bytes()

        docdocgo(
            sample_dataframe, sample_docx_with_marker, "{{TEST_MARKER}}", output_path
        )

        # Original file should be unchanged
        assert sample_docx_with_marker.read_bytes() == original_content

        # New file should exist and contain the table
        assert output_path.exists()
        doc = Document(output_path)
        assert len(doc.tables) == 1


class TestLargeDataFrames:
    """Test handling of large DataFrames and edge cases."""

    @pytest.fixture
    def data_output_dir(self):
        """Create tests/data directory for outputs."""
        data_dir = Path(__file__).parent / "data"
        data_dir.mkdir(exist_ok=True)
        return data_dir

    def test_empty_dataframe_error(self, data_output_dir):
        """Test that empty DataFrame raises appropriate error."""
        empty_df = pd.DataFrame()
        output_path = data_output_dir / "empty_dataframe_test.docx"

        with pytest.raises(ValueError, match="DataFrame cannot be empty"):
            create_document_with_table(empty_df, output_path)

    def test_large_dataframe_25_rows(self, data_output_dir):
        """Test creating document with 25 rows of data."""
        # Create DataFrame with 25 rows
        large_df = pd.DataFrame(
            {
                "ID": range(1, 26),
                "Name": [f"Person_{i:02d}" for i in range(1, 26)],
                "Age": [20 + (i % 40) for i in range(25)],
                "Score": [round(75 + (i * 2.5) % 25, 1) for i in range(25)],
                "Department": [f"Dept_{chr(65 + i % 5)}" for i in range(25)],
            }
        )

        output_path = data_output_dir / "large_dataframe_25_rows.docx"
        create_document_with_table(large_df, output_path)

        assert output_path.exists()

        # Verify document content
        doc = Document(output_path)
        tables = doc.tables
        assert len(tables) == 1

        table = tables[0]
        assert len(table.rows) == 26  # 1 header + 25 data rows
        assert len(table.columns) == 5

    def test_wide_dataframe_25_columns(self, data_output_dir):
        """Test creating document with 25 columns."""
        # Create DataFrame with 25 columns
        wide_data = {}
        for i in range(25):
            col_name = f"Col_{chr(65 + i % 26)}{i:02d}"
            wide_data[col_name] = [f"Value_{i}_{j}" for j in range(3)]

        wide_df = pd.DataFrame(wide_data)

        output_path = data_output_dir / "wide_dataframe_25_columns.docx"
        create_document_with_table(wide_df, output_path)

        assert output_path.exists()

        # Verify document content
        doc = Document(output_path)
        tables = doc.tables
        assert len(tables) == 1

        table = tables[0]
        assert len(table.rows) == 4  # 1 header + 3 data rows
        assert len(table.columns) == 25

    def test_dataframe_with_special_characters(self, data_output_dir):
        """Test DataFrame with special characters and Unicode."""
        special_df = pd.DataFrame(
            {
                "Unicode": ["café", "北京", "München", "🌟", "naïve"],
                "Special Chars": ["&<>\"'", "100%", "$1,000", "A&B", "x/y"],
                "Numbers": [1.5, -2.3, 0, 1234567890, 0.001],
            }
        )

        output_path = data_output_dir / "special_characters_test.docx"
        create_document_with_table(special_df, output_path)

        assert output_path.exists()

        # Verify document was created successfully
        doc = Document(output_path)
        assert len(doc.tables) == 1
