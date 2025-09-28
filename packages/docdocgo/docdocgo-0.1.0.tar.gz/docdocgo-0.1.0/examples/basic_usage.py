import pandas as pd
from docx import Document

import docdocgo

data = pd.DataFrame(
    {
        "Quarter": ["Q1", "Q2", "Q3", "Q4"],
        "Revenue": [120000, 135000, 142000, 158000],
        "Profit": [25000, 32000, 28000, 41000],
    }
)

doc = Document()
doc.add_heading("Quarterly Report", 0)
doc.add_paragraph("Here are the quarterly results:")
doc.add_paragraph("{{QUARTERLY_DATA}}")
doc.add_paragraph("Thank you for reviewing this report.")
doc.save("examples/template.docx")

docdocgo.insert_table(
    data, "examples/template.docx", "{{QUARTERLY_DATA}}", "examples/result.docx"
)
