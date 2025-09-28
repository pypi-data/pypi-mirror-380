"""
Generate a Word template with lorem ipsum content and docdocgo marker.

This script creates a professional Word document template with:
- Arial font throughout
- Professional headings and structure
- Lorem ipsum content
- {{QUARTERLY_PERFORMANCE_DATA}} marker for docdocgo replacement
"""

import argparse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


def generate_template(output_path="template.docx"):
    """Generate a Word template document."""
    # Create a new document
    doc = Document()

    # Set default style to Arial
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Inches(0.11)  # 11pt

    # Add title
    title = doc.add_heading("Business Report Template", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add subtitle
    subtitle = doc.add_heading("Quarterly Performance Analysis", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add some lorem ipsum content
    doc.add_paragraph(
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat."
    )

    doc.add_paragraph("")

    # Add section heading
    doc.add_heading("Executive Summary", level=2)

    doc.add_paragraph(
        "Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum."
    )

    doc.add_paragraph(
        "Sed ut perspiciatis unde omnis iste natus error sit voluptatem accusantium doloremque laudantium, totam rem aperiam, eaque ipsa quae ab illo inventore veritatis et quasi architecto beatae vitae dicta sunt explicabo."
    )

    doc.add_paragraph("")

    # Add quarterly data section
    doc.add_heading("Quarterly Financial Data", level=2)

    doc.add_paragraph("The following table contains our quarterly performance metrics:")

    doc.add_paragraph("")

    # Add the first marker for docdocgo
    marker_paragraph = doc.add_paragraph("{{QUARTERLY_DATA}}")
    marker_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # Add product data section
    doc.add_heading("Product Analysis", level=2)

    doc.add_paragraph("Below is our detailed product performance and inventory status:")

    doc.add_paragraph("")

    # Add the second marker for docdocgo
    marker_paragraph = doc.add_paragraph("{{PRODUCT_DATA}}")
    marker_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # Add conclusion section
    doc.add_heading("Conclusions", level=2)

    doc.add_paragraph(
        "Nemo enim ipsam voluptatem quia voluptas sit aspernatur aut odit aut fugit, sed quia consequuntur magni dolores eos qui ratione voluptatem sequi nesciunt."
    )

    doc.add_paragraph(
        "At vero eos et accusamus et iusto odio dignissimos ducimus qui blanditiis praesentibus voluptatum deleniti atque corrupti quos dolores et quas molestias excepturi sint occaecati cupiditate non provident."
    )

    doc.add_paragraph("")

    # Add footer
    footer_paragraph = doc.add_paragraph("---")
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer_paragraph = doc.add_paragraph(
        "Generated with docdocgo • Business Intelligence Team"
    )
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the template
    doc.save(output_path)
    print(f"✅ Template generated: {output_path}")
    print("📋 Template includes:")
    print("   • Arial font throughout")
    print("   • Professional headings and structure")
    print("   • Lorem ipsum content")
    print("   • {{QUARTERLY_DATA}} and {{PRODUCT_DATA}} markers for docdocgo")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Generate a Word template with docdocgo marker"
    )
    parser.add_argument(
        "--output",
        "-o",
        default="template.docx",
        help="Output file path (default: template.docx)",
    )
    args = parser.parse_args()

    generate_template(args.output)
