"""
Core functionality for docdocgo package.
"""

from pathlib import Path
from typing import Optional, Union

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT

from .styling import apply_table_style
from .utils import validate_dataframe, validate_file_path


def docdocgo(
    df: pd.DataFrame,
    path: Union[str, Path],
    marker: Optional[str] = None,
    output_path: Optional[Union[str, Path]] = None,
) -> None:
    """
    Convenience function to insert DataFrame into Word document.

    If marker is provided, replaces the marker in existing document.
    If marker is None, creates a new document with the table.

    Args:
        df: pandas DataFrame to insert
        path: Path to Word document (existing or new)
        marker: Marker string to replace (e.g., "{{DOCDOC_QUARTERLY_NUMBERS}}")
                If None, creates new document
        output_path: Optional output path. If None, overwrites the input file (for marker mode)
                    or uses the path parameter (for new document mode)
    """
    validate_dataframe(df)
    path = Path(path)

    if marker is not None:
        output = Path(output_path) if output_path else path
        insert_table(df, path, marker, output)
    else:
        output = Path(output_path) if output_path else path
        create_document_with_table(df, output)


def insert_table(
    df: pd.DataFrame,
    docx_path: Union[str, Path],
    marker: str,
    output_path: Optional[Union[str, Path]] = None,
) -> None:
    """
    Replace a marker in an existing Word document with a DataFrame table.

    Args:
        df: pandas DataFrame to insert
        docx_path: Path to existing Word document
        marker: Marker string to replace (e.g., "{{DOCDOC_QUARTERLY_NUMBERS}}")
        output_path: Optional output path. If None, overwrites the input document

    Raises:
        FileNotFoundError: If the Word document doesn't exist
        ValueError: If marker is not found in document
    """
    validate_dataframe(df)
    docx_path = Path(docx_path)
    validate_file_path(docx_path, must_exist=True)

    # Determine output path
    output = Path(output_path) if output_path else docx_path

    # Load the document
    doc = Document(docx_path)

    # Find and replace marker in paragraphs
    marker_found = False
    for paragraph in doc.paragraphs:
        if marker in paragraph.text:
            marker_found = True

            # Create table at end of document first
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Add headers
            header_cells = table.rows[0].cells
            for i, column in enumerate(df.columns):
                header_cells[i].text = str(column)

            # Add data rows
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            # Apply styling
            apply_table_style(table)

            # Replace the paragraph with the table
            tbl_element = table._tbl
            paragraph_element = paragraph._p
            paragraph_element.getparent().replace(paragraph_element, tbl_element)
            break

    if not marker_found:
        raise ValueError(f"Marker '{marker}' not found in document")

    # Save the document
    doc.save(output)


def create_document_with_table(df: pd.DataFrame, output_path: Union[str, Path]) -> None:
    """
    Create a new Word document with a DataFrame table.

    Args:
        df: pandas DataFrame to insert
        output_path: Path where the new Word document will be saved
    """
    validate_dataframe(df)
    output_path = Path(output_path)

    # Create new document
    doc = Document()

    # Add title
    doc.add_heading("Data Table", 0)

    # Create table
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Add headers
    header_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        header_cells[i].text = str(column)

    # Add data rows efficiently
    values = df.values
    for i in range(len(df)):
        row_cells = table.add_row().cells
        for j, value in enumerate(values[i]):
            row_cells[j].text = str(value)

    # Apply styling
    apply_table_style(table)

    # Save document
    doc.save(output_path)
