"""
Table styling functionality for docdocgo package.
"""

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor


def apply_table_style(table) -> None:
    """
    Apply professional styling to a Word table with dark blue headers.

    Args:
        table: python-docx Table object to style
    """
    # Set table style to a built-in table style
    table.style = "Table Grid"

    # Style header row
    header_row = table.rows[0]
    for cell in header_row.cells:
        # Set dark blue background for header
        cell_xml = cell._tc
        cell_properties = cell_xml.get_or_add_tcPr()

        # Add background color (dark blue)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "1f4e79")  # Dark blue color
        cell_properties.append(shading)

        # Set header text formatting
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)  # White text
        run.font.size = Inches(0.12)  # ~11pt

    # Style data rows
    for row_idx, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            # Alternate row colors for better readability
            if row_idx % 2 == 0:
                cell_xml = cell._tc
                cell_properties = cell_xml.get_or_add_tcPr()
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "f2f2f2")  # Light gray
                cell_properties.append(shading)

            # Set data text formatting
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Inches(0.11)  # ~10pt

    # Set column widths for better appearance
    for column in table.columns:
        for cell in column.cells:
            cell.width = Inches(1.5)

    # Set table alignment
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
