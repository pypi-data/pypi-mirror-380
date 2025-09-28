from __future__ import annotations

import os
import time
from typing import Iterable, List

from .base import DocumentLoader, make_document_id, normalize_path
from ..types import Document
from ..logging_utils import get_logger, log_operation


class TextFileLoader(DocumentLoader):
    def __init__(self):
        self.logger = get_logger()
    
    def load(self, paths: Iterable[str]) -> List[Document]:
        with log_operation("text_file_loading", loader="TextFileLoader"):
            documents: List[Document] = []
            processed_count = 0
            skipped_count = 0
            error_count = 0
            
            for path in paths:
                # Normalize path for cross-platform compatibility
                normalized_path = normalize_path(path)
                if not os.path.isfile(normalized_path):
                    self.logger.debug(f"Skipping non-file path: {path}")
                    skipped_count += 1
                    continue
                if not path.lower().endswith((".txt", ".md", ".rst")):
                    self.logger.debug(f"Skipping unsupported file type: {path}")
                    skipped_count += 1
                    continue
                try:
                    start_time = time.time()
                    with open(normalized_path, "r", encoding="utf-8", errors="ignore") as f:
                        text = f.read()
                    
                    document = Document(
                        id=make_document_id(path), 
                        path=normalized_path, 
                        text=text, 
                        meta={"type": "text", "size": len(text)}
                    )
                    documents.append(document)
                    
                    load_time = time.time() - start_time
                    self.logger.log_document_processing(
                        "loaded", 
                        normalized_path, 
                        f"size={len(text)} chars, time={load_time:.3f}s"
                    )
                    processed_count += 1
                    
                except Exception as e:
                    self.logger.log_error("text_file_loading", e, file_path=path)
                    error_count += 1
                    continue
            
            self.logger.debug(f"TextFileLoader completed: {processed_count} loaded, {skipped_count} skipped, {error_count} errors")
            return documents


class PDFFileLoader(DocumentLoader):
    def __init__(self) -> None:
        self.logger = get_logger()
        try:
            import fitz  # PyMuPDF  # type: ignore
            self._pymupdf_available = True
            self.logger.debug("PyMuPDF library available for PDF processing")
        except Exception as e:
            self._pymupdf_available = False
            self.logger.warning(f"PyMuPDF library not available: {e}")

    def load(self, paths: Iterable[str]) -> List[Document]:
        with log_operation("pdf_file_loading", loader="PDFFileLoader"):
            documents: List[Document] = []
            if not self._pymupdf_available:
                self.logger.warning("PDF processing skipped - PyMuPDF not available")
                return documents

            import fitz  # PyMuPDF  # type: ignore
            processed_count = 0
            skipped_count = 0
            error_count = 0

            for path in paths:
                # Normalize path for cross-platform compatibility
                normalized_path = normalize_path(path)
                if not os.path.isfile(normalized_path) or not path.lower().endswith(".pdf"):
                    self.logger.debug(f"Skipping non-PDF file: {path}")
                    skipped_count += 1
                    continue
                try:
                    start_time = time.time()
                    doc = fitz.open(normalized_path)
                    texts = []
                    page_count = doc.page_count
                    
                    for page_num in range(page_count):
                        try:
                            page = doc[page_num]
                            page_text = page.get_text() or ""
                            texts.append(page_text)
                        except Exception as e:
                            self.logger.debug(f"Error extracting text from page {page_num + 1}: {e}")
                            continue
                    
                    doc.close()
                    text = "\n".join(texts)
                    
                    document = Document(
                        id=make_document_id(path), 
                        path=normalized_path, 
                        text=text, 
                        meta={"type": "pdf", "pages": page_count, "size": len(text)}
                    )
                    documents.append(document)
                    
                    load_time = time.time() - start_time
                    self.logger.log_document_processing(
                        "loaded", 
                        normalized_path, 
                        f"pages={page_count}, size={len(text)} chars, time={load_time:.3f}s"
                    )
                    processed_count += 1
                    
                except Exception as e:
                    self.logger.log_error("pdf_file_loading", e, file_path=path)
                    error_count += 1
                    continue
            
            self.logger.debug(f"PDFFileLoader completed: {processed_count} loaded, {skipped_count} skipped, {error_count} errors")
            return documents


class DOCXFileLoader(DocumentLoader):
    def __init__(self) -> None:
        self.logger = get_logger()
        try:
            import docx  # type: ignore
            self._docx_available = True
            self.logger.debug("python-docx library available for DOCX processing")
        except Exception as e:
            self._docx_available = False
            self.logger.warning(f"python-docx library not available: {e}")

    def load(self, paths: Iterable[str]) -> List[Document]:
        with log_operation("docx_file_loading", loader="DOCXFileLoader"):
            documents: List[Document] = []
            if not self._docx_available:
                self.logger.warning("DOCX processing skipped - python-docx not available")
                return documents

            import docx  # type: ignore
            processed_count = 0
            skipped_count = 0
            error_count = 0

            for path in paths:
                # Normalize path for cross-platform compatibility
                normalized_path = normalize_path(path)
                if not os.path.isfile(normalized_path) or not path.lower().endswith(".docx"):
                    self.logger.debug(f"Skipping non-DOCX file: {path}")
                    skipped_count += 1
                    continue
                try:
                    start_time = time.time()
                    doc = docx.Document(normalized_path)
                    texts = []
                    paragraph_count = 0
                    
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            texts.append(paragraph.text.strip())
                            paragraph_count += 1
                    
                    text = "\n".join(texts)
                    
                    document = Document(
                        id=make_document_id(path), 
                        path=normalized_path, 
                        text=text, 
                        meta={"type": "docx", "paragraphs": paragraph_count, "size": len(text)}
                    )
                    documents.append(document)
                    
                    load_time = time.time() - start_time
                    self.logger.log_document_processing(
                        "loaded", 
                        normalized_path, 
                        f"paragraphs={paragraph_count}, size={len(text)} chars, time={load_time:.3f}s"
                    )
                    processed_count += 1
                    
                except Exception as e:
                    self.logger.log_error("docx_file_loading", e, file_path=path)
                    error_count += 1
                    continue
            
            self.logger.debug(f"DOCXFileLoader completed: {processed_count} loaded, {skipped_count} skipped, {error_count} errors")
            return documents


def load_documents(paths: Iterable[str]) -> List[Document]:
    logger = get_logger()
    
    with log_operation("document_loading", total_paths=len(list(paths)) if hasattr(paths, '__len__') else 'unknown'):
        loaders: List[DocumentLoader] = [TextFileLoader(), PDFFileLoader(), DOCXFileLoader()]
        all_docs: List[Document] = []
        
        logger.info(f"📄 Starting document loading with {len(loaders)} loaders")
        
        for i, loader in enumerate(loaders):
            loader_name = loader.__class__.__name__
            logger.debug(f"Running {loader_name} ({i+1}/{len(loaders)})")
            
            try:
                docs = loader.load(paths)
                all_docs.extend(docs)
                logger.debug(f"{loader_name} loaded {len(docs)} documents")
            except Exception as e:
                logger.log_error("document_loading", e, loader=loader_name)
                continue
        
        # Deduplicate by id
        seen = set()
        deduped: List[Document] = []
        duplicates = 0
        
        for d in all_docs:
            if d["id"] in seen:
                duplicates += 1
                logger.debug(f"Removing duplicate document: {d['id']}")
                continue
            seen.add(d["id"])
            deduped.append(d)
        
        logger.debug(f"Document loading completed: {len(deduped)} unique documents loaded, {duplicates} duplicates removed")
        return deduped


