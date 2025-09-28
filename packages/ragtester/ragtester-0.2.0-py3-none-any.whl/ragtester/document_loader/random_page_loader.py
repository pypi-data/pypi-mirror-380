from __future__ import annotations

import os
import random
from typing import Iterable, List, Set, Tuple
from dataclasses import dataclass

from .base import DocumentLoader, make_document_id
from ..types import Document


@dataclass
class PageSelection:
    """Represents a selected page from a document."""
    document_id: str
    page_number: int
    text: str
    word_count: int


class RandomPageLoader(DocumentLoader):
    """
    Document loader that implements random page selection method:
    1. Random selection of documents and pages
    2. Extract up to 300 words from selected page
    3. Track used pages to avoid repetition
    """
    
    def __init__(self, max_words_per_page: int = 300, seed: int = None):
        self.max_words_per_page = max_words_per_page
        self.used_selections: Set[Tuple[str, int]] = set()  # (document_id, page_number)
        self.available_selections: List[Tuple[str, int]] = []  # All possible (doc_id, page_num) pairs
        self.cached_pages: List[Tuple[str, int, str]] = []  # Cache for loaded pages (doc_id, page_num, page_text)
        self.document_paths_cache: List[str] = []  # Cache for document paths
        if seed is not None:
            random.seed(seed)
    
    def _extract_pages_from_pdf(self, path: str) -> List[str]:
        """Extract all pages from a PDF file."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF is required for PDF processing. Install with: pip install PyMuPDF")
        
        pages = []
        try:
            doc = fitz.open(path)
            # PDF loaded successfully
            for i in range(doc.page_count):
                try:
                    page = doc[i]
                    text = page.get_text() or ""
                    if text.strip():  # Only add non-empty pages
                        pages.append(text)
                    else:
                        pages.append("")  # Empty page if no text
                except Exception as e:
                    logger.debug(f"  Page {i+1}: extraction failed - {e}")
                    pages.append("")  # Empty page if extraction fails
            doc.close()
        except Exception as e:
            logger.error(f"Error reading PDF {path}: {e}")
            return []
        
        # Pages extracted successfully
        return pages
    
    def _count_words(self, text: str) -> int:
        """Count words in text."""
        return len(text.split())
    
    def _truncate_to_word_limit(self, text: str, max_words: int) -> str:
        """Truncate text to max_words, ensuring we don't cut words in half."""
        words = text.split()
        if len(words) <= max_words:
            return text
        return " ".join(words[:max_words])
    
    def _extract_text_from_docx(self, path: str) -> str:
        """Extract text from a DOCX file."""
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
        
        try:
            doc = docx.Document(path)
            texts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    texts.append(paragraph.text.strip())
            return "\n".join(texts)
        except Exception as e:
            logger.error(f"Error reading DOCX {path}: {e}")
            return ""

    def _load_document_pages(self, paths: Iterable[str]) -> List[Tuple[str, int, str]]:
        """
        Load all pages from all documents.
        Returns list of (document_id, page_number, page_text) tuples.
        """
        all_pages = []
        
        for path in paths:
            if not os.path.isfile(path):
                continue
                
            doc_id = make_document_id(path)
            
            if path.lower().endswith(".pdf"):
                pages = self._extract_pages_from_pdf(path)
                for page_num, page_text in enumerate(pages):
                    if page_text.strip():  # Only include non-empty pages
                        all_pages.append((doc_id, page_num, page_text))
            elif path.lower().endswith(".docx"):
                # Handle DOCX files properly
                text = self._extract_text_from_docx(path)
                if text.strip():
                    all_pages.append((doc_id, 0, text))
            else:
                # For other text files, treat the entire file as one "page"
                try:
                    with open(path, "r", encoding="utf-8", errors="ignore") as f:
                        text = f.read()
                    if text.strip():
                        all_pages.append((doc_id, 0, text))
                except Exception as e:
                    continue
        return all_pages
    
    def load(self, paths: Iterable[str]) -> List[Document]:
        """
        Load documents but don't return them yet - just prepare for random selection.
        The actual document creation happens in select_random_page().
        """
        # Load all available pages
        self.available_selections = []
        all_pages = self._load_document_pages(paths)
        
        if not all_pages:
            logger.warning(f"Warning: No pages loaded from documents: {list(paths)}")
            return []
        
        for doc_id, page_num, _ in all_pages:
            self.available_selections.append((doc_id, page_num))
        
        # Return empty list - documents will be created via select_random_page()
        return []
    
    def select_random_page(self, paths: Iterable[str]) -> PageSelection:
        """
        Select a random page that hasn't been used before.
        Returns a PageSelection object with the selected page content.
        """
        paths_list = list(paths)
        
        # Check if we need to reload (paths changed or no cache)
        if paths_list != self.document_paths_cache or not self.cached_pages:
            # print(f"Loading pages from {len(paths_list)} documents...")  # Commented out to reduce logging noise
            self.load(paths)
            # Cache the pages to avoid reloading
            self.cached_pages = self._load_document_pages(paths)
            self.document_paths_cache = paths_list.copy()
        
        # Find available (unused) selections
        available = [(doc_id, page_num) for doc_id, page_num in self.available_selections 
                    if (doc_id, page_num) not in self.used_selections]
        
        if not available:
            # If all pages have been used, reset and start over
            # print("All pages have been used, resetting selection pool...")  # Commented out to reduce logging noise
            self.used_selections.clear()
            available = self.available_selections.copy()
        
        # Check if we still have no available pages after reset
        if not available:
            logger.warning("Warning: Error selecting random page: Cannot choose from an empty sequence")
            raise ValueError("No pages available for selection")
        
        # Randomly select from available pages
        selected_doc_id, selected_page_num = random.choice(available)
        
        # Mark as used
        self.used_selections.add((selected_doc_id, selected_page_num))
        
        # Find the corresponding page text from cache
        selected_text = ""
        for doc_id, page_num, page_text in self.cached_pages:
            if doc_id == selected_doc_id and page_num == selected_page_num:
                selected_text = page_text
                break
        
        # Truncate to word limit
        truncated_text = self._truncate_to_word_limit(selected_text, self.max_words_per_page)
        word_count = self._count_words(truncated_text)
        
        return PageSelection(
            document_id=selected_doc_id,
            page_number=selected_page_num,
            text=truncated_text,
            word_count=word_count
        )
    
    def create_document_from_selection(self, selection: PageSelection, paths: Iterable[str]) -> Document:
        """Create a Document object from a PageSelection."""
        # Find the original path for this document
        original_path = ""
        for path in paths:
            if make_document_id(path) == selection.document_id:
                original_path = path
                break
        
        # Determine document type
        if original_path.lower().endswith(".pdf"):
            doc_type = "pdf_page"
        elif original_path.lower().endswith(".docx"):
            doc_type = "docx_page"
        else:
            doc_type = "text_page"
        
        return Document(
            id=f"{selection.document_id}_page_{selection.page_number}",
            path=original_path,
            text=selection.text,
            meta={
                "type": doc_type,
                "page_number": selection.page_number,
                "word_count": selection.word_count,
                "original_document_id": selection.document_id
            }
        )
    
    def get_remaining_pages_count(self) -> int:
        """Get the number of pages that haven't been used yet."""
        return len(self.available_selections) - len(self.used_selections)
    
    def reset_selections(self):
        """Reset the used selections to allow reusing pages."""
        self.used_selections.clear()
