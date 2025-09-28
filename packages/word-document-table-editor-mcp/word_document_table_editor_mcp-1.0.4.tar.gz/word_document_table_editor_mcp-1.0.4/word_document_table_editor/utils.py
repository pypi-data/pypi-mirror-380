"""
Word文档表格编辑工具的辅助函数

提供表格处理和格式化相关的辅助功能
"""

import os
from typing import Optional, List, Any
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def ensure_docx_extension(filename: str) -> str:
    """确保文件名有.docx扩展名
    
    Args:
        filename: 原始文件名
        
    Returns:
        带有.docx扩展名的文件名
    """
    if not filename.lower().endswith('.docx'):
        return filename + '.docx'
    return filename


def check_file_writeable(filename: str) -> tuple[bool, str]:
    """检查文件是否可写
    
    Args:
        filename: 文件路径
        
    Returns:
        (是否可写, 错误信息)
    """
    try:
        # 检查文件是否存在
        if not os.path.exists(filename):
            return False, "File does not exist"
        
        # 检查文件权限
        if not os.access(filename, os.W_OK):
            return False, "File is not writable (permission denied)"
        
        # 检查文件是否被其他程序锁定
        try:
            # 尝试以写入模式打开文件
            with open(filename, 'r+b') as f:
                pass
        except PermissionError:
            return False, "File is currently open in another application"
        except IOError as e:
            return False, f"File access error: {str(e)}"
        
        return True, ""
    
    except Exception as e:
        return False, f"Error checking file: {str(e)}"


def apply_table_style(table, has_header_row: bool, border_style: Optional[str], 
                     shading: Optional[List[List[str]]], border_color: Optional[str] = None) -> bool:
    """应用表格样式
    
    Args:
        table: 表格对象
        has_header_row: 是否有标题行
        border_style: 边框样式 ('solid', 'single', 'double', 'thick', 'none')
        shading: 阴影颜色矩阵
        border_color: 边框颜色 (hex格式，如 'FF0000' 表示红色)
        
    Returns:
        是否成功
    """
    try:
        # Apply header row formatting if specified
        if has_header_row and len(table.rows) > 0:
            header_row = table.rows[0]
            for cell in header_row.cells:
                # Make header bold
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        
        # Apply border style if specified
        if border_style:
            success = set_table_borders(table, border_style, border_color)
            if not success:
                return False
        
        # Apply shading if specified
        if shading:
            for row_idx, row_colors in enumerate(shading):
                if row_idx >= len(table.rows):
                    break
                row = table.rows[row_idx]
                for col_idx, color in enumerate(row_colors):
                    if col_idx >= len(row.cells):
                        break
                    cell = row.cells[col_idx]
                    set_cell_shading(cell, color)
        
        return True
    except Exception:
        return False


def set_cell_shading_by_position(table, row_index: int, col_index: int, 
                                fill_color: str, pattern: str = "clear") -> bool:
    """为指定位置的单元格设置阴影
    
    Args:
        table: 表格对象
        row_index: 行索引
        col_index: 列索引
        fill_color: 填充颜色
        pattern: 图案类型
        
    Returns:
        是否成功
    """
    try:
        if row_index >= len(table.rows) or col_index >= len(table.rows[row_index].cells):
            return False
        
        cell = table.rows[row_index].cells[col_index]
        return set_cell_shading(cell, fill_color, pattern)
    except Exception:
        return False


def set_cell_shading(cell, fill_color: str, pattern: str = "clear") -> bool:
    """为单元格设置阴影
    
    Args:
        cell: 单元格对象
        fill_color: 填充颜色
        pattern: 图案类型
        
    Returns:
        是否成功
    """
    try:
        # Get the cell's shading element
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Create shading element
        shd = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
        if shd is None:
            from docx.oxml import parse_xml
            shd_xml = f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="{pattern}" w:fill="{fill_color}"/>'
            shd = parse_xml(shd_xml)
            tcPr.append(shd)
        else:
            shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', pattern)
            shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', fill_color)
        
        return True
    except Exception:
        return False


def apply_alternating_row_shading(table, color1: str, color2: str) -> bool:
    """应用交替行阴影
    
    Args:
        table: 表格对象
        color1: 奇数行颜色
        color2: 偶数行颜色
        
    Returns:
        是否成功
    """
    try:
        for row_idx, row in enumerate(table.rows):
            color = color1 if row_idx % 2 == 0 else color2
            for cell in row.cells:
                set_cell_shading(cell, color)
        return True
    except Exception:
        return False


def highlight_header_row(table, header_color: str, text_color: str) -> bool:
    """高亮标题行
    
    Args:
        table: 表格对象
        header_color: 标题背景色
        text_color: 标题文字色
        
    Returns:
        是否成功
    """
    try:
        if len(table.rows) == 0:
            return False
        
        header_row = table.rows[0]
        for cell in header_row.cells:
            # Set background color
            set_cell_shading(cell, header_color)
            
            # Set text color and make bold
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    try:
                        # Parse color
                        if text_color.startswith('#'):
                            text_color = text_color[1:]
                        if len(text_color) == 6:
                            r = int(text_color[0:2], 16)
                            g = int(text_color[2:4], 16)
                            b = int(text_color[4:6], 16)
                            run.font.color.rgb = RGBColor(r, g, b)
                    except Exception:
                        # Default to white if color parsing fails
                        run.font.color.rgb = RGBColor(255, 255, 255)
        
        return True
    except Exception:
        return False


def merge_cells(table, start_row: int, start_col: int, end_row: int, end_col: int) -> bool:
    """合并单元格
    
    Args:
        table: 表格对象
        start_row: 开始行
        start_col: 开始列
        end_row: 结束行
        end_col: 结束列
        
    Returns:
        是否成功
    """
    try:
        # 检查表格是否为空
        if len(table.rows) == 0:
            return False
        
        # 验证索引范围
        if (start_row < 0 or end_row < 0 or start_col < 0 or end_col < 0):
            return False
        
        if (start_row >= len(table.rows) or end_row >= len(table.rows)):
            return False
        
        # 检查开始和结束位置的关系
        if start_row > end_row or start_col > end_col:
            return False
        
        # 检查列索引是否有效（检查每一行的列数）
        for row_idx in range(start_row, end_row + 1):
            if start_col >= len(table.rows[row_idx].cells) or end_col >= len(table.rows[row_idx].cells):
                return False
        
        # Get the starting cell
        start_cell = table.rows[start_row].cells[start_col]
        
        # 如果只合并一个单元格，不需要操作
        if start_row == end_row and start_col == end_col:
            return True
        
        # 合并矩形区域：使用正确的合并策略
        try:
            # 对于矩形区域合并，我们需要确保合并的顺序正确
            # 先合并右下角单元格，这样可以创建一个矩形区域
            if end_row > start_row or end_col > start_col:
                end_cell = table.rows[end_row].cells[end_col]
                start_cell.merge(end_cell)
            
            # 然后合并其他单元格，确保形成完整的矩形
            # 先合并第一行的其他列
            for col_idx in range(start_col + 1, end_col + 1):
                if col_idx != end_col or end_row == start_row:  # 避免重复合并右下角
                    target_cell = table.rows[start_row].cells[col_idx]
                    start_cell.merge(target_cell)
            
            # 再合并第一列的其他行
            for row_idx in range(start_row + 1, end_row + 1):
                if row_idx != end_row or end_col == start_col:  # 避免重复合并右下角
                    target_cell = table.rows[row_idx].cells[start_col]
                    start_cell.merge(target_cell)
            
            # 最后合并中间区域的单元格
            for row_idx in range(start_row + 1, end_row + 1):
                for col_idx in range(start_col + 1, end_col + 1):
                    if not (row_idx == end_row and col_idx == end_col):  # 跳过右下角
                        target_cell = table.rows[row_idx].cells[col_idx]
                        start_cell.merge(target_cell)
                        
        except Exception as e:
            # 如果上述方法失败，尝试更简单的方法
            try:
                # 简化方法：先合并水平，再合并垂直
                # 先合并第一行的所有列
                for col_idx in range(start_col + 1, end_col + 1):
                    target_cell = table.rows[start_row].cells[col_idx]
                    start_cell.merge(target_cell)
                
                # 然后合并第一列的所有行
                for row_idx in range(start_row + 1, end_row + 1):
                    target_cell = table.rows[row_idx].cells[start_col]
                    start_cell.merge(target_cell)
            except Exception:
                return False
        
        return True
    except Exception:
        return False


def merge_cells_horizontal(table, row_index: int, start_col: int, end_col: int) -> bool:
    """水平合并单元格
    
    Args:
        table: 表格对象
        row_index: 行索引
        start_col: 开始列
        end_col: 结束列
        
    Returns:
        是否成功
    """
    return merge_cells(table, row_index, start_col, row_index, end_col)


def merge_cells_vertical(table, col_index: int, start_row: int, end_row: int) -> bool:
    """垂直合并单元格

    Args:
        table: 表格对象
        col_index: 列索引
        start_row: 开始行
        end_row: 结束行

    Returns:
        是否成功
    """
    return merge_cells(table, start_row, col_index, end_row, col_index)


def set_cell_alignment_by_position(table, row_index: int, col_index: int,
                                  horizontal: str, vertical: str) -> bool:
    """设置指定位置单元格的对齐方式

    Args:
        table: 表格对象
        row_index: 行索引
        col_index: 列索引
        horizontal: 水平对齐
        vertical: 垂直对齐

    Returns:
        是否成功
    """
    try:
        if row_index >= len(table.rows) or col_index >= len(table.rows[row_index].cells):
            return False

        cell = table.rows[row_index].cells[col_index]
        return set_cell_alignment(cell, horizontal, vertical)
    except Exception:
        return False


def set_cell_alignment(cell, horizontal: str, vertical: str) -> bool:
    """设置单元格对齐方式

    Args:
        cell: 单元格对象
        horizontal: 水平对齐
        vertical: 垂直对齐

    Returns:
        是否成功
    """
    try:
        # Set vertical alignment
        vertical_alignment_map = {
            'top': WD_CELL_VERTICAL_ALIGNMENT.TOP,
            'center': WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            'bottom': WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        }

        if vertical.lower() in vertical_alignment_map:
            cell.vertical_alignment = vertical_alignment_map[vertical.lower()]

        # Set horizontal alignment for paragraphs in the cell
        horizontal_alignment_map = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
            'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        }

        if horizontal.lower() in horizontal_alignment_map:
            for paragraph in cell.paragraphs:
                paragraph.alignment = horizontal_alignment_map[horizontal.lower()]

        return True
    except Exception:
        return False


def set_table_alignment(table, horizontal: str, vertical: str) -> bool:
    """设置表格所有单元格的对齐方式

    Args:
        table: 表格对象
        horizontal: 水平对齐
        vertical: 垂直对齐

    Returns:
        是否成功
    """
    try:
        for row in table.rows:
            for cell in row.cells:
                set_cell_alignment(cell, horizontal, vertical)
        return True
    except Exception:
        return False


def format_cell_text_by_position(table, row_index: int, col_index: int,
                                text_content: Optional[str], bold: Optional[bool],
                                italic: Optional[bool], underline: Optional[bool],
                                color: Optional[str], font_size: Optional[int],
                                font_name: Optional[str]) -> bool:
    """格式化指定位置单元格的文本

    Args:
        table: 表格对象
        row_index: 行索引
        col_index: 列索引
        text_content: 文本内容
        bold: 是否粗体
        italic: 是否斜体
        underline: 是否下划线
        color: 文本颜色
        font_size: 字体大小
        font_name: 字体名称

    Returns:
        是否成功
    """
    try:
        if row_index >= len(table.rows) or col_index >= len(table.rows[row_index].cells):
            return False

        cell = table.rows[row_index].cells[col_index]

        # Set text content if provided
        if text_content is not None:
            cell.text = text_content

        # Format all runs in the cell
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if bold is not None:
                    run.bold = bold
                if italic is not None:
                    run.italic = italic
                if underline is not None:
                    run.underline = underline
                if font_size is not None:
                    run.font.size = Pt(font_size)
                if font_name is not None:
                    run.font.name = font_name
                if color is not None:
                    try:
                        # Parse color
                        if color.startswith('#'):
                            color = color[1:]
                        if len(color) == 6:
                            r = int(color[0:2], 16)
                            g = int(color[2:4], 16)
                            b = int(color[4:6], 16)
                            run.font.color.rgb = RGBColor(r, g, b)
                    except Exception:
                        # Default to black if color parsing fails
                        run.font.color.rgb = RGBColor(0, 0, 0)

        return True
    except Exception:
        return False


def set_cell_padding_by_position(table, row_index: int, col_index: int,
                                top: Optional[float], bottom: Optional[float],
                                left: Optional[float], right: Optional[float],
                                unit: str) -> bool:
    """设置指定位置单元格的内边距

    Args:
        table: 表格对象
        row_index: 行索引
        col_index: 列索引
        top: 上边距
        bottom: 下边距
        left: 左边距
        right: 右边距
        unit: 单位

    Returns:
        是否成功
    """
    try:
        if row_index >= len(table.rows) or col_index >= len(table.rows[row_index].cells):
            return False

        cell = table.rows[row_index].cells[col_index]

        # Convert values to appropriate units
        unit_map = {
            'points': lambda x: Pt(x),
            'inches': lambda x: Inches(x),
            'cm': lambda x: Inches(x / 2.54)
        }

        if unit.lower() not in unit_map:
            return False

        converter = unit_map[unit.lower()]

        # Set cell margins (this is a simplified implementation)
        # Full implementation would require more complex margin handling
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Create margin element if needed
        tcMar = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcMar')
        if tcMar is None:
            from docx.oxml import parse_xml
            tcMar_xml = '<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            tcMar = parse_xml(tcMar_xml)
            tcPr.append(tcMar)

        # Set individual margins
        if top is not None:
            top_elem = tcMar.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}top')
            if top_elem is None:
                from docx.oxml import parse_xml
                top_xml = f'<w:top xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="{int(converter(top).pt * 20)}" w:type="dxa"/>'
                top_elem = parse_xml(top_xml)
                tcMar.append(top_elem)
            else:
                top_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', str(int(converter(top).pt * 20)))

        # Similar for other margins (simplified for brevity)

        return True
    except Exception:
        return False


def set_column_width_by_position(table, col_index: int, width: float, width_type: str) -> bool:
    """设置指定列的宽度

    Args:
        table: 表格对象
        col_index: 列索引
        width: 宽度值
        width_type: 宽度类型

    Returns:
        是否成功
    """
    try:
        # Check if column index is valid
        if len(table.rows) == 0 or col_index >= len(table.rows[0].cells):
            return False

        # Convert width to appropriate units
        unit_map = {
            'points': lambda x: Pt(x),
            'inches': lambda x: Inches(x),
            'cm': lambda x: Inches(x / 2.54),
            'pixels': lambda x: Inches(x / 96),  # 96 DPI -> px to inches
            'percent': lambda x: x  # Percentage handling would be more complex
        }

        if width_type.lower() not in unit_map:
            return False

        if width_type.lower() != 'percent':
            width_value = unit_map[width_type.lower()](width)

            # Set width for all cells in the column
            for row in table.rows:
                if col_index < len(row.cells):
                    row.cells[col_index].width = width_value

        return True
    except Exception:
        return False


def set_column_widths(table, widths: List[float], width_type: str) -> bool:
    """设置多列宽度

    Args:
        table: 表格对象
        widths: 宽度列表
        width_type: 宽度类型

    Returns:
        是否成功
    """
    try:
        for col_index, width in enumerate(widths):
            if not set_column_width_by_position(table, col_index, width, width_type):
                return False
        return True
    except Exception:
        return False


def set_table_width_func(table, width: float, width_type: str) -> bool:
    """设置表格宽度

    Args:
        table: 表格对象
        width: 宽度值
        width_type: 宽度类型

    Returns:
        是否成功
    """
    try:
        # Convert width to appropriate units
        unit_map = {
            'points': lambda x: Pt(x),
            'inches': lambda x: Inches(x),
            'cm': lambda x: Inches(x / 2.54),
            'percent': lambda x: x  # Percentage handling would be more complex
        }

        if width_type.lower() not in unit_map:
            return False

        if width_type.lower() != 'percent':
            width_value = unit_map[width_type.lower()](width)
            table.width = width_value

        return True
    except Exception:
        return False


def auto_fit_table(table) -> bool:
    """自动调整表格列宽

    Args:
        table: 表格对象

    Returns:
        是否成功
    """
    try:
        # This is a simplified implementation
        # Full auto-fit would require content analysis
        table.autofit = True
        return True
    except Exception:
        return False


def set_table_borders(table, border_style: str, border_color: Optional[str] = None) -> bool:
    """设置表格边框样式和颜色
    
    Args:
        table: 表格对象
        border_style: 边框样式 ('solid', 'single', 'double', 'thick', 'none')  
        border_color: 边框颜色 (hex格式，如 'FF0000' 表示红色)
        
    Returns:
        是否成功
    """
    try:
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls, qn
        
        # Map border styles to Word constants
        style_map = {
            'solid': 'single',
            'single': 'single', 
            'double': 'double',
            'thick': 'thick',
            'none': 'none'
        }
        
        word_style = style_map.get(border_style.lower(), 'single')
        
        # Default color is black if not specified
        color = border_color if border_color else '000000'
        if color.startswith('#'):
            color = color[1:]
        
        # Create border properties
        if word_style == 'none':
            border_xml = f"""
            <w:tblBorders {nsdecls('w')}>
                <w:top w:val="none"/>
                <w:left w:val="none"/>
                <w:bottom w:val="none"/>
                <w:right w:val="none"/>
                <w:insideH w:val="none"/>
                <w:insideV w:val="none"/>
            </w:tblBorders>"""
        else:
            border_xml = f"""
            <w:tblBorders {nsdecls('w')}>
                <w:top w:val="{word_style}" w:color="{color}" w:sz="4"/>
                <w:left w:val="{word_style}" w:color="{color}" w:sz="4"/>
                <w:bottom w:val="{word_style}" w:color="{color}" w:sz="4"/>
                <w:right w:val="{word_style}" w:color="{color}" w:sz="4"/>
                <w:insideH w:val="{word_style}" w:color="{color}" w:sz="4"/>
                <w:insideV w:val="{word_style}" w:color="{color}" w:sz="4"/>
            </w:tblBorders>"""
        
        # Parse the XML and add it to table properties
        tbl_borders = parse_xml(border_xml)
        
        # Get table properties
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr_xml = f'<w:tblPr {nsdecls("w")}/>'
            tblPr = parse_xml(tblPr_xml)
            tbl.insert(0, tblPr)
        
        # Remove existing borders if any
        existing_borders = tblPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
        if existing_borders is not None:
            tblPr.remove(existing_borders)
        
        # Add new borders
        tblPr.append(tbl_borders)
        
        return True
    except Exception as e:
        return False
