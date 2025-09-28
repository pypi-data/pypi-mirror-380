"""
Word文档表格编辑工具函数

提供十五个核心的文档表格编辑功能
"""

import os
from typing import Optional, List
from .utils import (
    ensure_docx_extension,
    check_file_writeable,
    apply_table_style,
    set_cell_shading_by_position,
    apply_alternating_row_shading,
    highlight_header_row,
    merge_cells,
    merge_cells_horizontal,
    merge_cells_vertical,
    set_cell_alignment_by_position,
    set_table_alignment,
    format_cell_text_by_position,
    set_cell_padding_by_position,
    set_column_width_by_position,
    set_column_widths,
    set_table_width_func,
    auto_fit_table
)


async def format_table(filename: str, table_index: int, 
                      has_header_row: Optional[bool] = None,
                      border_style: Optional[str] = None,
                      shading: Optional[List[List[str]]] = None,
                      border_color: Optional[str] = None) -> str:
    """Format a table with borders, shading, and structure.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        has_header_row: If True, formats the first row as a header
        border_style: Style for borders ('none', 'single', 'double', 'thick', 'solid')
        shading: 2D list of cell background colors (by row and column)
        border_color: Border color (hex string like 'FF0000' for red)
    """
    from docx import Document
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")
    
    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1}).")
        table = doc.tables[table_index]
        success = apply_table_style(table, has_header_row or False, border_style, shading, border_color)
        if success:
            doc.save(filename)
            return f"Table at index {table_index} formatted successfully."
        else:
            raise RuntimeError(f"Failed to format table at index {table_index}.")
    except Exception as e:
        raise RuntimeError(f"Failed to format table: {str(e)}")


async def set_table_cell_shading(filename: str, table_index: int, row_index: int, col_index: int, 
                                fill_color: str, pattern: str = "clear") -> str:
    """Apply shading/filling to a specific table cell.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        row_index: Row index of the cell (0-based)
        col_index: Column index of the cell (0-based)
        fill_color: Background color (hex string like "FF0000" or "red")
        pattern: Shading pattern ("clear", "solid", "pct10", "pct20", etc.)
    """
    from docx import Document
    
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        row_index = int(row_index)
        col_index = int(col_index)
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: table_index, row_index, and col_index must be integers")

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")
    
    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1}).")
        table = doc.tables[table_index]
        success = set_cell_shading_by_position(table, row_index, col_index, fill_color, pattern)
        if success:
            doc.save(filename)
            return f"Cell shading applied successfully to table {table_index}, cell ({row_index},{col_index})."
        else:
            raise RuntimeError("Failed to apply cell shading. Check that indices are valid.")
    except Exception as e:
        raise RuntimeError(f"Failed to apply cell shading: {str(e)}")


async def apply_table_alternating_rows(filename: str, table_index: int, 
                                     color1: str = "FFFFFF", color2: str = "F2F2F2") -> str:
    """Apply alternating row colors to a table for better readability.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        color1: Color for odd rows (hex string, default white)
        color2: Color for even rows (hex string, default light gray)
    """
    from docx import Document
    
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: table_index must be an integer")
    
    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")
    
    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1}).")
        table = doc.tables[table_index]
        success = apply_alternating_row_shading(table, color1, color2)
        if success:
            doc.save(filename)
            return f"Alternating row shading applied successfully to table {table_index}."
        else:
            raise RuntimeError("Failed to apply alternating row shading.")
    except Exception as e:
        raise RuntimeError(f"Failed to apply alternating row shading: {str(e)}")


async def highlight_table_header(filename: str, table_index: int, 
                               header_color: str = "4472C4", text_color: str = "FFFFFF") -> str:
    """Apply special highlighting to table header row.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        header_color: Background color for header (hex string, default blue)
        text_color: Text color for header (hex string, default white)
    """
    from docx import Document
    
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: table_index must be an integer")
    
    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")
    
    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1}).")
        table = doc.tables[table_index]
        success = highlight_header_row(table, header_color, text_color)
        if success:
            doc.save(filename)
            return f"Header highlighting applied successfully to table {table_index}."
        else:
            raise RuntimeError("Failed to apply header highlighting.")
    except Exception as e:
        raise RuntimeError(f"Failed to apply header highlighting: {str(e)}")


async def merge_table_cells(filename: str, table_index: int, start_row: int, start_col: int,
                           end_row: int, end_col: int) -> str:
    """Merge cells in a rectangular area of a table.

    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        start_row: Starting row index (0-based)
        start_col: Starting column index (0-based)
        end_row: Ending row index (0-based, inclusive)
        end_col: Ending column index (0-based, inclusive)
    """
    from docx import Document

    filename = ensure_docx_extension(filename)

    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        start_row = int(start_row)
        start_col = int(start_col)
        end_row = int(end_row)
        end_col = int(end_col)
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: all indices must be integers")

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")

    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1}).")
        table = doc.tables[table_index]
        
        # 提供更详细的验证信息
        if len(table.rows) == 0:
            raise ValueError("Table is empty and cannot be merged.")
        
        if start_row < 0 or end_row < 0 or start_col < 0 or end_col < 0:
            raise ValueError("All indices must be non-negative.")
        
        if start_row >= len(table.rows) or end_row >= len(table.rows):
            raise ValueError(f"Row indices out of range. Table has {len(table.rows)} rows (0-{len(table.rows)-1}).")
        
        if start_row > end_row or start_col > end_col:
            raise ValueError("Start position must be less than or equal to end position.")
        
        # 检查列索引
        for row_idx in range(start_row, end_row + 1):
            row_cell_count = len(table.rows[row_idx].cells)
            if start_col >= row_cell_count or end_col >= row_cell_count:
                raise ValueError(f"Column indices out of range. Row {row_idx} has {row_cell_count} columns (0-{row_cell_count-1}).")
        
        success = merge_cells(table, start_row, start_col, end_row, end_col)
        if success:
            doc.save(filename)
            return f"Cells merged successfully in table {table_index} from ({start_row},{start_col}) to ({end_row},{end_col})."
        else:
            raise RuntimeError("Failed to merge cells due to an internal error.")
    except Exception as e:
        raise RuntimeError(f"Failed to merge cells: {str(e)}")


async def merge_table_cells_horizontal(filename: str, table_index: int, row_index: int,
                                      start_col: int, end_col: int) -> str:
    """Merge cells horizontally in a single row.

    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        row_index: Row index (0-based)
        start_col: Starting column index (0-based)
        end_col: Ending column index (0-based, inclusive)
    """
    from docx import Document

    filename = ensure_docx_extension(filename)

    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        row_index = int(row_index)
        start_col = int(start_col)
        end_col = int(end_col)
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: all indices must be integers")

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")

    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1}).")
        table = doc.tables[table_index]
        success = merge_cells_horizontal(table, row_index, start_col, end_col)
        if success:
            doc.save(filename)
            return f"Cells merged horizontally in table {table_index}, row {row_index}, columns {start_col}-{end_col}."
        else:
            raise RuntimeError("Failed to merge cells horizontally. Check that indices are valid.")
    except Exception as e:
        raise RuntimeError(f"Failed to merge cells horizontally: {str(e)}")


async def merge_table_cells_vertical(filename: str, table_index: int, col_index: int,
                                    start_row: int, end_row: int) -> str:
    """Merge cells vertically in a single column.

    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        col_index: Column index (0-based)
        start_row: Starting row index (0-based)
        end_row: Ending row index (0-based, inclusive)
    """
    from docx import Document

    filename = ensure_docx_extension(filename)

    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        col_index = int(col_index)
        start_row = int(start_row)
        end_row = int(end_row)
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: all indices must be integers")

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")

    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1}).")
        table = doc.tables[table_index]
        success = merge_cells_vertical(table, col_index, start_row, end_row)
        if success:
            doc.save(filename)
            return f"Cells merged vertically in table {table_index}, column {col_index}, rows {start_row}-{end_row}."
        else:
            raise RuntimeError("Failed to merge cells vertically. Check that indices are valid.")
    except Exception as e:
        raise RuntimeError(f"Failed to merge cells vertically: {str(e)}")


async def set_table_cell_alignment(filename: str, table_index: int, row_index: int, col_index: int,
                                  horizontal: str = "left", vertical: str = "top") -> str:
    """Set text alignment for a specific table cell.

    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        row_index: Row index of the cell (0-based)
        col_index: Column index of the cell (0-based)
        horizontal: Horizontal alignment ("left", "center", "right", "justify")
        vertical: Vertical alignment ("top", "center", "bottom")
    """
    from docx import Document

    filename = ensure_docx_extension(filename)

    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        row_index = int(row_index)
        col_index = int(col_index)
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: table_index, row_index, and col_index must be integers")

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")

    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1}).")
        table = doc.tables[table_index]
        success = set_cell_alignment_by_position(table, row_index, col_index, horizontal, vertical)
        if success:
            doc.save(filename)
            return f"Cell alignment set successfully for table {table_index}, cell ({row_index},{col_index})."
        else:
            raise RuntimeError("Failed to set cell alignment. Check that indices are valid.")
    except Exception as e:
        raise RuntimeError(f"Failed to set cell alignment: {str(e)}")


async def set_table_alignment_all(filename: str, table_index: int,
                                 horizontal: str = "left", vertical: str = "top") -> str:
    """Set text alignment for all cells in a table.

    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        horizontal: Horizontal alignment ("left", "center", "right", "justify")
        vertical: Vertical alignment ("top", "center", "bottom")
    """
    from docx import Document

    filename = ensure_docx_extension(filename)

    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: table_index must be an integer")

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")

    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1}).")
        table = doc.tables[table_index]
        success = set_table_alignment(table, horizontal, vertical)
        if success:
            doc.save(filename)
            return f"Alignment set successfully for all cells in table {table_index}."
        else:
            raise RuntimeError("Failed to set table alignment.")
    except Exception as e:
        raise RuntimeError(f"Failed to set table alignment: {str(e)}")


async def format_table_cell_text(filename: str, table_index: int, row_index: int, col_index: int,
                                text_content: Optional[str] = None, bold: Optional[bool] = None,
                                italic: Optional[bool] = None, underline: Optional[bool] = None,
                                color: Optional[str] = None, font_size: Optional[int] = None,
                                font_name: Optional[str] = None) -> str:
    """Format text within a specific table cell.

    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        row_index: Row index of the cell (0-based)
        col_index: Column index of the cell (0-based)
        text_content: New text content for the cell (optional)
        bold: Set text bold (True/False)
        italic: Set text italic (True/False)
        underline: Set text underlined (True/False)
        color: Text color (e.g., 'red', 'blue', hex code)
        font_size: Font size in points
        font_name: Font name/family
    """
    from docx import Document

    filename = ensure_docx_extension(filename)

    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        row_index = int(row_index)
        col_index = int(col_index)
        if font_size is not None:
            font_size = int(font_size)
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: table_index, row_index, col_index, and font_size must be integers")

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")

    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1}).")
        table = doc.tables[table_index]
        success = format_cell_text_by_position(table, row_index, col_index, text_content,
                                              bold, italic, underline, color, font_size, font_name)
        if success:
            doc.save(filename)
            return f"Cell text formatted successfully for table {table_index}, cell ({row_index},{col_index})."
        else:
            raise RuntimeError("Failed to format cell text. Check that indices are valid.")
    except Exception as e:
        raise RuntimeError(f"Failed to format cell text: {str(e)}")


async def set_table_cell_padding(filename: str, table_index: int, row_index: int, col_index: int,
                                top: Optional[float] = None, bottom: Optional[float] = None,
                                left: Optional[float] = None, right: Optional[float] = None,
                                unit: str = "points") -> str:
    """Set padding/margins for a specific table cell.

    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        row_index: Row index of the cell (0-based)
        col_index: Column index of the cell (0-based)
        top: Top padding value
        bottom: Bottom padding value
        left: Left padding value
        right: Right padding value
        unit: Unit of measurement ("points", "inches", "cm")
    """
    from docx import Document

    filename = ensure_docx_extension(filename)

    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        row_index = int(row_index)
        col_index = int(col_index)
        if top is not None:
            top = float(top)
        if bottom is not None:
            bottom = float(bottom)
        if left is not None:
            left = float(left)
        if right is not None:
            right = float(right)
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: indices must be integers, padding values must be numbers")

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")

    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1}).")
        table = doc.tables[table_index]
        success = set_cell_padding_by_position(table, row_index, col_index, top, bottom, left, right, unit)
        if success:
            doc.save(filename)
            return f"Cell padding set successfully for table {table_index}, cell ({row_index},{col_index})."
        else:
            raise RuntimeError("Failed to set cell padding. Check that indices are valid.")
    except Exception as e:
        raise RuntimeError(f"Failed to set cell padding: {str(e)}")


async def set_table_column_width(filename: str, table_index: int, col_index: int,
                                width: float, width_type: str = "points") -> str:
    """Set the width of a specific table column.

    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        col_index: Column index (0-based)
        width: Width value
        width_type: Unit of measurement ("points", "inches", "cm", "percent")
    """
    from docx import Document

    filename = ensure_docx_extension(filename)

    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        col_index = int(col_index)
        width = float(width)
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: indices must be integers, width must be a number")

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")

    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1}).")
        table = doc.tables[table_index]
        success = set_column_width_by_position(table, col_index, width, width_type)
        if success:
            doc.save(filename)
            return f"Column width set successfully for table {table_index}, column {col_index}."
        else:
            raise RuntimeError("Failed to set column width. Check that column index is valid.")
    except Exception as e:
        raise RuntimeError(f"Failed to set column width: {str(e)}")


async def set_table_column_widths(filename: str, table_index: int, widths: List[float] | List[dict],
                                 width_type: str = "points") -> str:
    """Set the widths of multiple table columns.

    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        widths: List of width values for each column
        width_type: Unit of measurement ("points", "inches", "cm", "percent")
    """
    from docx import Document

    filename = ensure_docx_extension(filename)

    # Ensure numeric parameters are the correct type, and support object form
    try:
        table_index = int(table_index)
        # allow two formats:
        # 1) [number, number, ...]
        # 2) [{"col_index": int, "width": number, "width_type": str?}, ...]
        parsed_widths: List[float] | List[dict] = widths
        is_object_list = isinstance(widths, list) and len(widths) > 0 and isinstance(widths[0], dict)
        if is_object_list:
            # We'll apply selectively per object below
            parsed_widths = widths  # type: ignore
        else:
            # treat as simple list of numbers
            parsed_widths = [float(w) for w in widths]  # type: ignore
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: table_index must be integer, widths must be a list of numbers or objects")

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")

    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1}).")
        table = doc.tables[table_index]
        # If object list, set per specified column; otherwise set sequentially
        if isinstance(parsed_widths, list) and len(parsed_widths) > 0 and isinstance(parsed_widths[0], dict):
            overall_success = True
            for item in parsed_widths:  # type: ignore
                try:
                    col_index = int(item.get('col_index'))
                    col_width = float(item.get('width'))
                    col_width_type = (item.get('width_type') or width_type)
                    if not set_column_width_by_position(table, col_index, col_width, col_width_type):
                        overall_success = False
                except Exception:
                    overall_success = False
            success = overall_success
        else:
            success = set_column_widths(table, parsed_widths, width_type)  # type: ignore
        if success:
            doc.save(filename)
            return f"Column widths set successfully for table {table_index}."
        else:
            raise RuntimeError("Failed to set column widths. Check that the number of widths matches the table columns.")
    except Exception as e:
        raise RuntimeError(f"Failed to set column widths: {str(e)}")


async def set_table_width(filename: str, table_index: int, width: float,
                         width_type: str = "points") -> str:
    """Set the overall width of a table.

    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        width: Width value
        width_type: Unit of measurement ("points", "inches", "cm", "percent")
    """
    from docx import Document

    filename = ensure_docx_extension(filename)

    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        width = float(width)
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: table_index must be integer, width must be a number")

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")

    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1}).")
        table = doc.tables[table_index]
        success = set_table_width_func(table, width, width_type)
        if success:
            doc.save(filename)
            return f"Table width set successfully for table {table_index}."
        else:
            raise RuntimeError("Failed to set table width.")
    except Exception as e:
        raise RuntimeError(f"Failed to set table width: {str(e)}")


async def auto_fit_table_columns(filename: str, table_index: int) -> str:
    """Set table columns to auto-fit their content.

    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
    """
    from docx import Document

    filename = ensure_docx_extension(filename)

    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: table_index must be an integer")

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")

    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1}).")
        table = doc.tables[table_index]
        success = auto_fit_table(table)
        if success:
            doc.save(filename)
            return f"Auto-fit applied successfully to table {table_index}."
        else:
            raise RuntimeError("Failed to apply auto-fit to table.")
    except Exception as e:
        raise RuntimeError(f"Failed to auto-fit table columns: {str(e)}")
