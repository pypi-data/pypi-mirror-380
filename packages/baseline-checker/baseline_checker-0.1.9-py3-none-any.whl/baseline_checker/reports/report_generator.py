import json
import csv
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt
from jinja2 import Template
from fpdf import FPDF
from colorama import Fore, Style
# ---------------- JSON ----------------
def save_json(data, filename="baseline_report.json"):
    with open(filename, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4)
    print(f"{Fore.YELLOW}[INFO]{Style.RESET_ALL} scanned report saved as ({filename})")

# ---------------- CSV ----------------
def save_csv(data, filename="baseline_report.csv"):
    with open(filename, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Folder", "Files Scanned", "Baseline Features", "Non-Baseline Features"])
        for folder in data.get("folders", []):
            writer.writerow([
                folder["name"],
                folder["files_scanned"],
                "; ".join(folder["baseline_features"]),
                "; ".join(folder["non_baseline_features"])
            ])
    print(f"{Fore.YELLOW}[INFO]{Style.RESET_ALL} scanned report saved as ({filename})")

# ---------------- DOCX ----------------


def save_word(data, filename="baseline_report.docx"):
    doc = Document()
    
    # Title
    title = doc.add_heading("Baseline Checker Report", level=0)
    title.alignment = 1  # Centered
    
    # Summary
    doc.add_paragraph(f"Total Files Scanned: {data.get('total_files_scanned',0)}")
    doc.add_paragraph(f"Total Baseline Features Used: {len(data.get('baseline_features', []))}")
    doc.add_paragraph(f"Total Non-Baseline Features Used: {len(data.get('non_baseline_features', []))}")
    
    # Per folder report
    for folder in data.get("folders", []):
        folder_title = doc.add_heading(folder["name"], level=1)
        
        doc.add_paragraph(f"Files Scanned: {folder['files_scanned']}")
        
        # Baseline features in green
        p = doc.add_paragraph()
        p.add_run("Baseline Features: ").bold = True
        for feat in folder['baseline_features']:
            r = p.add_run(feat + ", ")
            r.font.color.rgb = RGBColor(0, 128, 0)
            r.font.size = Pt(11)
        
        # Non-Baseline features in red
        p = doc.add_paragraph()
        p.add_run("Non-Baseline Features: ").bold = True
        for feat in folder['non_baseline_features']:
            r = p.add_run(feat + ", ")
            r.font.color.rgb = RGBColor(255, 0, 0)
            r.font.size = Pt(11)
        
        doc.add_paragraph("")  # Empty line for spacing
    doc.save(filename)
    print(f"{Fore.YELLOW}[INFO]{Style.RESET_ALL} scanned report saved as ({filename})")


# ---------------- PDF ----------------


def save_pdf(data, filename="baseline_report.pdf"):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # Title
    pdf.set_font("Arial", "B", 18)
    pdf.set_text_color(0, 0, 128)  # Dark Blue
    pdf.cell(0, 12, "Baseline Checker Report", ln=True, align="C")
    pdf.ln(5)
    
    # Summary
    pdf.set_font("Arial", "B", 12)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 8, f"Total Files Scanned: {data.get('total_files_scanned', 0)}", ln=True)
    pdf.cell(0, 8, f"Total Baseline Features Used: {len(data.get('baseline_features', []))}", ln=True)
    pdf.cell(0, 8, f"Total Non-Baseline Features Used: {len(data.get('non_baseline_features', []))}", ln=True)
    pdf.ln(5)
    
    # Per folder report
    for folder in data.get("folders", []):
        pdf.set_font("Arial", "B", 14)
        pdf.set_text_color(0, 0, 128)
        pdf.cell(0, 8, folder["name"], ln=True)
        
        pdf.set_font("Arial", "", 12)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 8, f"Files Scanned: {folder['files_scanned']}", ln=True)
        
        # Baseline features in green
        pdf.set_text_color(0, 128, 0)
        baseline_text = ', '.join(folder['baseline_features']) if folder['baseline_features'] else 'None'
        pdf.multi_cell(0, 8, f"Baseline Features: {baseline_text}")
        
        # Non-Baseline features in red
        pdf.set_text_color(255, 0, 0)
        nonbaseline_text = ', '.join(folder['non_baseline_features']) if folder['non_baseline_features'] else 'None'
        pdf.multi_cell(0, 8, f"Non-Baseline Features: {nonbaseline_text}")
        
        pdf.ln(3)
        pdf.set_text_color(0, 0, 0)  # Reset color for next folder
    
    pdf.output(filename)
    print(f"{Fore.YELLOW}[INFO]{Style.RESET_ALL} scanned report saved as ({filename})")


# ---------------- HTML ----------------
def save_html(data, filename="baseline_report.html"):
    template_html = """
    <html>
    <head>
        <title>Baseline Checker Report</title>
        <style>
            body { font-family: Arial, sans-serif; }
            h1 { color: #2E8B57; }
            h2 { color: #4682B4; }
            .baseline { color: green; }
            .non-baseline { color: red; }
            table { border-collapse: collapse; width: 100%; margin-bottom: 20px; }
            th, td { border: 1px solid #ccc; padding: 8px; text-align: left; }
            th { background-color: #f2f2f2; }
        </style>
    </head>
    <body>
        <h1>Baseline Checker Report</h1>
        <p><strong>Total Files Scanned:</strong> {{ total_files_scanned }}</p>
        <p><strong>Total Baseline Features Used:</strong> {{ baseline_count }}</p>
        <p><strong>Total Non-Baseline Features Used:</strong> {{ non_baseline_count }}</p>
        {% for folder in folders %}
        <h2>{{ folder.name }}</h2>
        <table>
            <tr><th>Files Scanned</th><td>{{ folder.files_scanned }}</td></tr>
            <tr><th>Baseline Features</th><td>{{ folder.baseline_features | join(', ') if folder.baseline_features else 'None' }}</td></tr>
            <tr><th>Non-Baseline Features</th><td>{{ folder.non_baseline_features | join(', ') if folder.non_baseline_features else 'None' }}</td></tr>
        </table>
        {% endfor %}
    </body>
    </html>
    """
    template = Template(template_html)
    html_content = template.render(
        total_files_scanned=data.get("total_files_scanned", 0),
        baseline_count=len(data.get("baseline_features", [])),
        non_baseline_count=len(data.get("non_baseline_features", [])),
        folders=data.get("folders", [])
    )
    Path(filename).write_text(html_content, encoding="utf-8")
    print(f"{Fore.YELLOW}[INFO]{Style.RESET_ALL} scanned report saved as ({filename})")